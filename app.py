#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""WordGrab · 桌面 GUI（pywebview）
左侧历史 + 文稿点读联动 + 播放器 + 说话人改名 + 搜索 + 导出
"""
import os, sys, json, time, uuid, shutil, threading, datetime, re, subprocess, hashlib, sqlite3
import ai_service
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

HERE = os.path.dirname(os.path.abspath(__file__))
UI = os.path.join(HERE, "ui")
DEFAULT_DATA = os.path.join(HERE, "data")
STORAGE_CONFIG = os.path.join(HERE, "storage.json")


def configured_data_directory():
    """读取独立于数据目录的启动指针，目录搬迁后重启仍能找到数据。"""
    try:
        with open(STORAGE_CONFIG, encoding="utf-8") as file:
            configured = json.load(file).get("data_directory")
        if configured:
            return os.path.abspath(os.path.expanduser(str(configured)))
    except (OSError, ValueError, TypeError, AttributeError):
        pass
    return DEFAULT_DATA


DATA = configured_data_directory()
try:
    os.makedirs(DATA, exist_ok=True)
except OSError as exc:
    print(f"[storage] 无法使用已设置的数据目录 {DATA!r}: {exc}; 已临时使用默认目录", flush=True)
    DATA = DEFAULT_DATA
    os.makedirs(DATA, exist_ok=True)
INDEX = os.path.join(DATA, "index.json")
SETTINGS = os.path.join(DATA, "settings.json")
TRASH = os.path.join(DATA, ".trash")
INDEX_DB = os.path.join(DATA, "index.db")
APP_VERSION = "1.2.0"

THEME_KEYS = {
    "aurora-sea", "solar-bloom", "lavender-haze", "tide-ember",
    "midnight-prism", "matcha-mist", "rose-quartz", "sandstone-glow",
    "deep-ocean", "graphite-pearl", "rainbow-glow",
}

DEFAULT_SETTINGS = {
    "theme": "aurora-sea",
    "reopen_last": True,
    "auto_open_import": True,
    "default_speed": 1.0,
    "skip_seconds": 15,
    "auto_diarization": True,
    "transcription_mode": "accuracy",
    "export_format": "txt",
    "export_directory": os.path.expanduser("~/Documents"),
    "filename_rule": "source_date",
    "font_size": "standard",
    "list_density": "standard",
    "appearance": "light",
    "follow_system": False,
    "delete_audio_with_transcript": True,
    "last_item_id": "",
    "ai_base_url": "",
    "ai_model": "",
    "ai_summary_template": "general",
    "ai_privacy_host": "",
}

SETTING_ENUMS = {
    "transcription_mode": {"speed", "accuracy"},
    "export_format": {"docx", "pdf", "txt"},
    "filename_rule": {"source", "source_date"},
    "font_size": {"small", "standard", "large"},
    "list_density": {"compact", "standard"},
    "appearance": {"system", "light", "dark"},
}

MODEL_CACHE_MARKERS = ("paraformer", "fsmn_vad", "punc_ct", "campplus")

AUDIO_EXT = {".m4a": "audio/mp4", ".mp3": "audio/mpeg", ".wav": "audio/wav",
             ".aac": "audio/aac", ".mp4": "video/mp4", ".mov": "video/quicktime",
             ".m4v": "video/mp4", ".flac": "audio/flac", ".ogg": "audio/ogg"}

# 供 HTTP 处理器（拖拽上传）回调到 Api 实例
API_REF = None
INDEX_LOCK = threading.RLock()
MAX_UPLOAD_BYTES = 2 * 1024 * 1024 * 1024  # 2 GB，避免误拖入超大视频占满磁盘

# 转写进度共享给前端轮询（后台线程只写这个 dict，不从子线程调 evaluate_js —— 后者在 macOS 上易崩）
PROGRESS = {}                      # iid -> {stage, pct, info, status, msg, title, partial}
AI_TASKS = {}                      # task_id -> AI 总结任务状态
AI_TASKS_LOCK = threading.RLock()
TRANSCRIBE_LOCK = threading.Lock()  # 序列化转写，避免并发共用同一个模型实例出错
DELETED = set()                    # 转写期间被用户删除的 iid，两阶段落盘前都要检查
_DRAG_STRIP_CLASS = None           # macOS 原生透明拖动带（延迟创建，避免非 macOS 导入 AppKit）
_DRAG_STRIPS = {}                  # NSWindow id -> drag view，持有引用避免被回收
_DRAG_OBSERVERS = {}               # NSWindow id -> 通知监听状态，缩放/全屏后重新定位拖动带

# ---------- 历史存储 ----------

def atomic_write_json(path, value):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    temporary = path + ".tmp"
    with open(temporary, "w", encoding="utf-8") as file:
        json.dump(value, file, ensure_ascii=False, indent=2)
        file.flush()
        os.fsync(file.fileno())
    os.replace(temporary, path)

def load_index():
    with INDEX_LOCK:
        _ensure_index_db()
        try:
            with sqlite3.connect(INDEX_DB) as connection:
                rows = connection.execute("SELECT payload FROM items ORDER BY position").fetchall()
            return [json.loads(row[0]) for row in rows]
        except (OSError, sqlite3.Error, ValueError):
            pass
        if os.path.exists(INDEX):
            try:
                with open(INDEX, encoding="utf-8") as f:
                    value = json.load(f)
                return value if isinstance(value, list) else []
            except (OSError, ValueError):
                backup = INDEX + ".bak"
                try:
                    with open(backup, encoding="utf-8") as f:
                        value = json.load(f)
                    return value if isinstance(value, list) else []
                except (OSError, ValueError):
                    return []
    return []


def save_index(items):
    with INDEX_LOCK:
        _ensure_index_db()
        payloads = [(index, json.dumps(item, ensure_ascii=False)) for index, item in enumerate(items)]
        with sqlite3.connect(INDEX_DB) as connection:
            connection.execute("BEGIN")
            connection.execute("DELETE FROM items")
            connection.executemany("INSERT INTO items(position, payload) VALUES (?, ?)", payloads)
            connection.commit()
        temporary = INDEX + ".tmp"
        backup = INDEX + ".bak"
        with open(temporary, "w", encoding="utf-8") as f:
            json.dump(items, f, ensure_ascii=False, indent=2)
            f.flush()
            os.fsync(f.fileno())
        if os.path.exists(INDEX):
            shutil.copy2(INDEX, backup)
        os.replace(temporary, INDEX)


def _ensure_index_db():
    """首次运行把旧 index.json 迁移到 SQLite；JSON 继续保留作人工恢复备份。"""
    global INDEX_DB
    if os.path.exists(INDEX_DB):
        return
    os.makedirs(os.path.dirname(INDEX_DB), exist_ok=True)
    with sqlite3.connect(INDEX_DB) as connection:
        connection.execute("CREATE TABLE IF NOT EXISTS items (position INTEGER PRIMARY KEY, payload TEXT NOT NULL)")
        if os.path.exists(INDEX):
            try:
                with open(INDEX, encoding="utf-8") as file:
                    old_items = json.load(file)
                connection.executemany(
                    "INSERT INTO items(position, payload) VALUES (?, ?)",
                    [(index, json.dumps(item, ensure_ascii=False)) for index, item in enumerate(old_items or [])],
                )
            except (OSError, ValueError, TypeError):
                pass
        connection.commit()


def normalize_settings(data=None):
    raw = data if isinstance(data, dict) else {}
    out = dict(DEFAULT_SETTINGS)
    for key in out:
        if key in raw:
            out[key] = raw[key]
    if out["theme"] not in THEME_KEYS:
        out["theme"] = DEFAULT_SETTINGS["theme"]
    for key, allowed in SETTING_ENUMS.items():
        if out[key] not in allowed:
            out[key] = DEFAULT_SETTINGS[key]
    if "appearance" not in raw:
        out["appearance"] = "system" if raw.get("follow_system") else DEFAULT_SETTINGS["appearance"]
    out["follow_system"] = out["appearance"] == "system"
    if out["skip_seconds"] not in {5, 10, 15, 30}:
        out["skip_seconds"] = 15
    try:
        out["default_speed"] = float(out["default_speed"])
    except (TypeError, ValueError):
        out["default_speed"] = 1.0
    if out["default_speed"] not in {0.25, 0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0}:
        out["default_speed"] = 1.0
    for key in ("reopen_last", "auto_open_import", "auto_diarization",
                "delete_audio_with_transcript"):
        out[key] = bool(out[key])
    out["last_item_id"] = str(out.get("last_item_id") or "")
    out["ai_base_url"] = str(out.get("ai_base_url") or "").strip().rstrip("/")
    out["ai_model"] = str(out.get("ai_model") or "").strip()
    out["ai_privacy_host"] = str(out.get("ai_privacy_host") or "")
    directory = os.path.expanduser(str(out.get("export_directory") or ""))
    out["export_directory"] = directory if os.path.isdir(directory) else os.path.expanduser("~/Documents")
    return out


def load_settings():
    if os.path.exists(SETTINGS):
        try:
            with open(SETTINGS, encoding="utf-8") as f:
                return normalize_settings(json.load(f))
        except (OSError, ValueError, TypeError):
            pass
    return normalize_settings()


def save_settings(data):
    data = normalize_settings(data)
    with open(SETTINGS, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


def apply_settings_patch(patch):
    current = load_settings()
    if isinstance(patch, dict):
        for key, value in patch.items():
            if key in DEFAULT_SETTINGS:
                current[key] = value
    return save_settings(current)


def path_size(path):
    total = 0
    if not os.path.exists(path):
        return 0
    for root, _, files in os.walk(path):
        for name in files:
            try:
                total += os.path.getsize(os.path.join(root, name))
            except OSError:
                pass
    return total


def _set_data_globals(directory):
    global DATA, INDEX, SETTINGS, TRASH, INDEX_DB
    DATA = os.path.abspath(os.path.expanduser(directory))
    INDEX = os.path.join(DATA, "index.json")
    SETTINGS = os.path.join(DATA, "settings.json")
    TRASH = os.path.join(DATA, ".trash")
    INDEX_DB = os.path.join(DATA, "index.db")


def _save_data_directory_pointer(directory):
    payload = {"data_directory": os.path.abspath(os.path.expanduser(directory))}
    temporary = STORAGE_CONFIG + ".tmp"
    with open(temporary, "w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)
    os.replace(temporary, STORAGE_CONFIG)


def migrate_data_directory(directory):
    """把现有数据复制到空目录，确认成功后再切换并删除旧目录。"""
    source = os.path.abspath(DATA)
    raw_target = str(directory or "").strip()
    if not raw_target:
        return {"ok": False, "message": "没有选择新的保存位置"}
    target = os.path.abspath(os.path.expanduser(raw_target))
    if source == target:
        return {"ok": True, "data_path": source, "moved": False,
                "message": "当前已经使用这个位置"}
    try:
        common = os.path.commonpath([source, target])
    except ValueError:
        common = ""
    if common in {source, target}:
        return {"ok": False, "message": "新位置不能包含当前数据目录，请选择其他空文件夹"}
    if any(value.get("status") in {"running", "queued"} for value in PROGRESS.values()):
        return {"ok": False, "message": "正在转写音频，请等待完成后再更改保存位置"}

    try:
        os.makedirs(target, exist_ok=True)
        meaningful_entries = [
            name for name in os.listdir(target)
            if name not in {".DS_Store", ".localized"}
        ]
        if meaningful_entries:
            return {"ok": False, "message": "为避免覆盖文件，请选择一个空文件夹"}
        probe = os.path.join(target, f".wordgrab-write-test-{uuid.uuid4().hex}")
        with open(probe, "w", encoding="utf-8") as file:
            file.write("ok")
        os.remove(probe)
    except OSError as exc:
        return {"ok": False, "message": f"无法使用这个文件夹：{exc}"}

    copied = []
    try:
        for name in os.listdir(source):
            if name in {".DS_Store", ".localized"}:
                continue
            source_entry = os.path.join(source, name)
            target_entry = os.path.join(target, name)
            copied.append(target_entry)
            if os.path.isdir(source_entry) and not os.path.islink(source_entry):
                shutil.copytree(source_entry, target_entry)
            else:
                shutil.copy2(source_entry, target_entry)
        _save_data_directory_pointer(target)
        _set_data_globals(target)
    except Exception as exc:
        for copied_path in reversed(copied):
            try:
                if os.path.isdir(copied_path) and not os.path.islink(copied_path):
                    shutil.rmtree(copied_path)
                elif os.path.lexists(copied_path):
                    os.remove(copied_path)
            except OSError:
                pass
        return {"ok": False, "message": f"移动数据失败，原位置未改变：{exc}"}

    warning = ""
    try:
        shutil.rmtree(source)
    except OSError:
        warning = "新位置已生效；旧文件夹暂时无法删除，可稍后手动清理。"
    return {
        "ok": True,
        "data_path": DATA,
        "moved": True,
        "message": warning or "文稿和录音已移动到新位置",
    }


def model_cache_dirs():
    roots = [
        os.path.expanduser("~/.cache/modelscope/models/iic"),
        os.path.expanduser("~/.cache/modelscope/hub/models/iic"),
    ]
    found = []
    for root in roots:
        if not os.path.isdir(root):
            continue
        for name in os.listdir(root):
            full = os.path.join(root, name)
            if os.path.isdir(full) and any(marker in name.lower() for marker in MODEL_CACHE_MARKERS):
                found.append(full)
    return found


def safe_filename(name):
    cleaned = re.sub(r'[\\/:*?"<>|]+', "-", str(name or "转写")).strip(" .")
    return cleaned[:120] or "转写"


def preserve_item_audio(iid, title):
    source = audio_path_of(iid)
    if not source or not os.path.isfile(source):
        return None
    target_dir = os.path.join(DATA, "保留的录音")
    os.makedirs(target_dir, exist_ok=True)
    ext = os.path.splitext(source)[1]
    base = safe_filename(title)
    target = os.path.join(target_dir, base + ext)
    if os.path.exists(target):
        stamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
        target = os.path.join(target_dir, f"{base}-{stamp}{ext}")
    shutil.move(source, target)
    return target


EXPORT_COLORS = (
    "#FF5C4D", "#FF9F1C", "#FFD23F", "#2EC4B6",
    "#3A86FF", "#8338EC", "#FF4D8D", "#10B981",
)
SPEAKER_COLOR_POOL = (
    "#FF675B", "#FF9F1C", "#F59E0B", "#FFD23F",
    "#84CC16", "#10B981", "#14B8A6", "#06B6D4",
    "#3A86FF", "#6366F1", "#8338EC", "#A855F7",
    "#D946EF", "#EC4899", "#F43F5E", "#EF4444",
)
EXPORT_BLOCK_MAX_CHARS = 320


def format_export_time(milliseconds):
    seconds = max(0, int(milliseconds or 0) // 1000)
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def split_export_text(text, limit=EXPORT_BLOCK_MAX_CHARS):
    """优先在中文标点处切块，防止一个说话人段落高于整页。"""
    remaining = str(text or "")
    chunks = []
    while len(remaining) > limit:
        cut = -1
        for mark in "。！？；.!?;\n":
            candidate = remaining.rfind(mark, int(limit * .58), limit)
            cut = max(cut, candidate)
        cut = cut + 1 if cut >= 0 else limit
        chunks.append(remaining[:cut])
        remaining = remaining[cut:]
    if remaining or not chunks:
        chunks.append(remaining)
    return chunks


def build_export_rows(segments, speakers, duration_seconds, speaker_colors=None):
    speaker_colors = speaker_colors or {}
    prepared = []
    raw = list(segments or [])
    duration_ms = max(0, int(float(duration_seconds or 0) * 1000))
    for index, segment in enumerate(raw):
        text = str(segment.get("text") or "")
        if not text:
            continue
        speaker_index = int(segment.get("spk", 0) or 0)
        start_ms = max(0, int(segment.get("start", 0) or 0))
        explicit_end = int(segment.get("end", 0) or 0)
        next_start = (int(raw[index + 1].get("start", 0) or 0)
                      if index + 1 < len(raw) else duration_ms)
        end_ms = explicit_end if explicit_end > start_ms else max(start_ms, next_start)
        chunks = split_export_text(text)
        consumed = 0
        text_length = max(1, len(text))
        for chunk in chunks:
            chunk_start = start_ms + round((end_ms - start_ms) * consumed / text_length)
            consumed += len(chunk)
            chunk_end = start_ms + round((end_ms - start_ms) * consumed / text_length)
            prepared.append({
                "speaker_index": speaker_index,
                "start_ms": chunk_start,
                "end_ms": max(chunk_start, chunk_end),
                "text": chunk,
            })

    merged = []
    for row in prepared:
        previous = merged[-1] if merged else None
        can_merge = (
            previous
            and previous["speaker_index"] == row["speaker_index"]
            and len(previous["text"]) + len(row["text"]) <= EXPORT_BLOCK_MAX_CHARS
        )
        if can_merge:
            joiner = " " if (previous["text"][-1:].isascii()
                              and previous["text"][-1:].isalnum()
                              and row["text"][:1].isascii()
                              and row["text"][:1].isalnum()) else ""
            previous["text"] += joiner + row["text"]
            previous["end_ms"] = row["end_ms"]
        else:
            merged.append(dict(row))

    for row in merged:
        index = row["speaker_index"]
        row["speaker"] = speakers.get(str(index), f"说话人 {index + 1}")
        row["start_time"] = format_export_time(row["start_ms"])
        row["end_time"] = format_export_time(row["end_ms"])
        row["time"] = row["start_time"]
        row["color"] = speaker_colors.get(str(index), EXPORT_COLORS[index % len(EXPORT_COLORS)])
    return merged


def export_payload(iid):
    data = load_item(iid)
    meta = next((item for item in load_index() if item["id"] == iid), {})
    speakers = data.get("speakers", {})
    duration = meta.get("duration") or 0
    speaker_colors, colors_changed = ensure_speaker_colors(iid, data)
    if colors_changed:
        save_item(iid, data)
    rows = build_export_rows(data.get("segments", []), speakers, duration, speaker_colors)
    participant_names = []
    for row in rows:
        if row["speaker"] not in participant_names:
            participant_names.append(row["speaker"])
    if not participant_names:
        participant_names = list(speakers.values()) or ["说话人 1"]
    participant_count = meta.get("n_speakers") or len(participant_names) or 1
    joined_participants = " · ".join(participant_names)
    if len(joined_participants) <= 11:
        participants_text = joined_participants
    elif participant_count > 2:
        participants_text = f"{participant_names[0][:8]} 等{participant_count}人"
    elif len(participant_names) > 1:
        participants_text = f"{participant_names[0][:6]} · {participant_names[1][:6]}"
    else:
        participants_text = participant_names[0][:10] + "…"

    created = str(meta.get("created") or "").replace("T", " ").strip()
    created_parts = created.split()
    recorded_text = (f"{created_parts[0]} · {created_parts[1][:5]}"
                     if len(created_parts) >= 2 else created)
    word_count = sum(1 for row in rows for char in row["text"] if not char.isspace())
    return {
        "title": meta.get("title") or "转写",
        "created": created,
        "recorded_text": recorded_text or "—",
        "duration": duration,
        "duration_text": format_export_time(float(duration) * 1000),
        "word_count": word_count,
        "word_count_text": f"{word_count:,} 字",
        "speaker_count": participant_count,
        "participants_text": participants_text,
        "rows": rows,
    }


def ai_summary_export_payload(iid):
    base = export_payload(iid)
    source_title = base["title"]
    summary = ai_service.load_summary(item_dir(iid))
    if not summary or not isinstance(summary.get("result"), dict):
        raise ValueError("当前文稿还没有可导出的 AI 总结")
    result = summary["result"]
    data = load_item(iid)
    segments = data.get("segments") or []
    colors = EXPORT_COLORS

    def source_time(value):
        ids = []
        def collect(child):
            if isinstance(child, dict):
                raw = child.get("source_segment_ids")
                if isinstance(raw, list): ids.extend(raw)
                for nested in child.values(): collect(nested)
            elif isinstance(child, list):
                for nested in child: collect(nested)
        collect(value)
        for source in ids:
            match = re.fullmatch(r"seg-(\d+)", str(source))
            if match and int(match.group(1)) < len(segments):
                return format_export_time(segments[int(match.group(1))].get("start") or 0)
        return "00:00"

    def item_text(item):
        if isinstance(item, str): return item.strip()
        if not isinstance(item, dict): return ""
        text = str(item.get("text") or item.get("summary") or item.get("quote")
                   or item.get("title") or item.get("task") or "").strip()
        details = []
        if item.get("owner"): details.append("负责人：" + str(item["owner"]))
        if item.get("due") or item.get("deadline"): details.append("截止：" + str(item.get("due") or item.get("deadline")))
        if item.get("status"): details.append("状态：" + str(item["status"]))
        when = item.get("time")
        if when: details.append("时间：" + str(when))
        return text + (("（" + "；".join(details) + "）") if details else "")

    sections = []
    def add(label, value):
        if value in (None, "", []): return
        if isinstance(value, list):
            text = "\n".join(f"• {item_text(item)}" for item in value if item_text(item))
        elif isinstance(value, dict):
            text = "\n".join(f"{key}：{item_text(item)}" for key, item in value.items() if item_text(item))
        else: text = str(value).strip()
        if text: sections.append((label, text, source_time(value)))

    if summary.get("template") == "meeting":
        add("会议目的", result.get("purpose")); add("讨论主题", result.get("topics"))
        add("关键结论", result.get("conclusions")); add("已确认决策", result.get("decisions"))
        add("待办事项", result.get("actions")); add("风险与待确认", result.get("risks"))
    else:
        overview = result.get("overview") or {}
        if isinstance(overview, dict):
            type_names = {"meeting":"会议","interview":"访谈","lecture":"讲座","call":"通话","memo":"备忘","other":"其他"}
            overview_text = "\n".join(filter(None, [
                "分析标题：" + str(overview.get("title") or ""),
                "内容类型：" + type_names.get(overview.get("type"), str(overview.get("type") or "其他")),
                "时长：" + str(overview.get("duration") or base["duration_text"]),
                "参与人：" + "、".join(str(x) for x in (overview.get("speakers") or [])),
            ]))
            add("概览", overview_text)
        else: add("概览", overview)
        add("一句话总结", result.get("one_line_summary"))
        add("智能摘要", result.get("summary"))
        add("章节脉络", result.get("chapters"))
        add("关键信息 / 要点", result.get("key_points"))
        add("待办与行动项", result.get("action_items"))
        add("关键词", "、".join(str(x) for x in (result.get("keywords") or [])))
        add("亮点 / 金句", result.get("highlights"))
        decisions = result.get("decisions") or {}
        if isinstance(decisions, dict):
            add("已确认决策", decisions.get("decided")); add("分歧", decisions.get("disagreements")); add("待解问题", decisions.get("open"))
        add("后续建议", result.get("suggestions"))

    rows = []
    for index, (label, text, start_time) in enumerate(sections):
        rows.append({"speaker_index": index, "speaker": label, "start_ms": 0, "end_ms": 0,
                     "start_time": start_time, "end_time": start_time, "time": start_time,
                     "text": text, "color": colors[index % len(colors)]})
    analysis_title = ((result.get("overview") or {}).get("title")
                      if isinstance(result.get("overview"), dict) else "")
    base["title"] = f"{base['title']} · {analysis_title or 'AI 分析'}"
    base["participants_text"] = "AI 内容分析"
    base["rows"] = rows
    base["word_count"] = sum(1 for row in rows for char in row["text"] if not char.isspace())
    base["word_count_text"] = f"{base['word_count']:,} 字"
    base["is_ai_analysis"] = True
    base["source_title"] = source_title
    base["analysis_title"] = analysis_title or "AI 内容分析"
    base["template_name"] = summary.get("template_name") or ("会议纪要" if summary.get("template") == "meeting" else "通用摘要")
    base["ai_result"] = result
    base["ai_template"] = summary.get("template") or "general"
    return base


def write_txt_export(path, payload):
    lines = [payload["title"], payload["created"], ""]
    for row in payload["rows"]:
        lines.append(
            f"[{row['start_time']} - {row['end_time']}] {row['speaker']}：{row['text']}"
        )
    with open(path, "w", encoding="utf-8") as file:
        file.write("\n".join(lines) + "\n")


def write_docx_export(path, payload):
    import io
    import struct
    import zlib

    from docx import Document
    from docx.enum.table import (WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE,
                                 WD_TABLE_ALIGNMENT)
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    body_font = "Heiti SC"
    brand_font = "Arial Black"
    page_width_pt, left_margin_pt, right_margin_pt = 595.276, 40.0, 41.0
    content_width_pt = page_width_pt - left_margin_pt - right_margin_pt

    def set_run_font(run, name=body_font, size=12, color="191919", bold=False):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)
        r_fonts.set(qn("w:eastAsia"), name)

    def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for tag, value in (("top", top), ("start", start),
                           ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{tag}"))
            if node is None:
                node = OxmlElement(f"w:{tag}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(int(round(value * 20))))
            node.set(qn("w:type"), "dxa")

    def set_cell_borders(cell, **edges):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge, options in edges.items():
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            for key, value in options.items():
                node.set(qn(f"w:{key}"), str(value))

    def set_cell_fill(cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.first_child_found_in("w:shd")
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), color.lstrip("#"))

    def set_table_geometry(table, widths_pt, indent_pt=0):
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(round(sum(widths_pt) * 20))))
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_ind.set(qn("w:w"), str(int(round(indent_pt * 20))))
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths_pt:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(round(width * 20))))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths_pt):
                tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    cell._tc.get_or_add_tcPr().append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(round(width * 20))))

    def add_exact_spacer(height_pt):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(height_pt)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("\u200b")
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(255, 255, 255)

    def clear_table_borders(cell):
        nil = {"val": "nil"}
        set_cell_borders(cell, top=nil, left=nil, bottom=nil, right=nil,
                         insideH=nil, insideV=nil)

    def color_strip_image(color_values=EXPORT_COLORS, width=1400, height=28, gap=3):
        """用内存 PNG 固定色条宽度，避免不同 Word 引擎压缩空表格。"""
        colors_rgb = [
            tuple(bytes.fromhex(color.lstrip("#"))) for color in color_values
        ]
        usable = width - gap * (len(colors_rgb) - 1)
        boundaries = [round(index * usable / len(colors_rgb))
                      for index in range(len(colors_rgb) + 1)]
        pixels = []
        for index, color in enumerate(colors_rgb):
            pixels.extend([color] * (boundaries[index + 1] - boundaries[index]))
            if index < len(colors_rgb) - 1:
                pixels.extend([(255, 255, 255)] * gap)
        row = b"\x00" + b"".join(bytes(pixel) for pixel in pixels)
        raw = row * height

        def chunk(kind, data):
            return (struct.pack(">I", len(data)) + kind + data
                    + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF))

        stream = io.BytesIO(
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", zlib.compress(raw, 9))
            + chunk(b"IEND", b"")
        )
        stream.seek(0)
        return stream

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Pt(left_margin_pt)
    section.right_margin = Pt(right_margin_pt)
    section.top_margin = Pt(14.84)
    section.bottom_margin = Pt(40)
    section.footer_distance = Pt(9)
    document.core_properties.title = payload["title"]
    document.core_properties.subject = "WordGrab 本地语音转录稿"
    document.core_properties.author = "WordGrab"

    styles = document.styles
    styles["Normal"].font.name = body_font
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    styles["Normal"].font.size = Pt(14)
    normal_format = styles["Normal"].paragraph_format
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)

    header = document.add_table(rows=1, cols=2)
    set_table_geometry(header, [47.2, content_width_pt - 47.2])
    header.rows[0].height = Pt(40)
    header.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    icon_cell, brand_cell = header.rows[0].cells
    for cell in (icon_cell, brand_cell):
        clear_table_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(icon_cell)
    set_cell_margins(brand_cell)
    icon_paragraph = icon_cell.paragraphs[0]
    icon_paragraph.paragraph_format.space_after = Pt(0)
    icon_paragraph.add_run().add_picture(
        os.path.join(UI, "icon_1024.png"), width=Pt(40), height=Pt(40))
    brand_paragraph = brand_cell.paragraphs[0]
    brand_paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(brand_paragraph.add_run("WordGrab"), brand_font, 20, "191919")

    add_exact_spacer(9.68)
    rainbow = document.add_table(rows=1, cols=1)
    set_table_geometry(rainbow, [content_width_pt])
    rainbow.rows[0].height = Pt(10)
    rainbow.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    rainbow_cell = rainbow.cell(0, 0)
    set_cell_margins(rainbow_cell)
    clear_table_borders(rainbow_cell)
    rainbow_paragraph = rainbow_cell.paragraphs[0]
    rainbow_paragraph.paragraph_format.space_before = Pt(0)
    rainbow_paragraph.paragraph_format.space_after = Pt(0)
    rainbow_paragraph.add_run().add_picture(
        color_strip_image(), width=Pt(content_width_pt), height=Pt(10))

    add_exact_spacer(16.87)
    metadata = document.add_table(rows=2, cols=3)
    metadata_width = content_width_pt - 1
    metadata_col = metadata_width / 3
    set_table_geometry(metadata, [metadata_col] * 3, indent_pt=1)
    labels = ("录制时间", "时长 · 字数", "参会人")
    values = (payload["recorded_text"],
              f"{payload['duration_text']} · {payload['word_count_text']}",
              payload["participants_text"])
    for row_index, row in enumerate(metadata.rows):
        row.height = Pt(20)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell_index, cell in enumerate(row.cells):
            set_cell_margins(cell, top=1.5, start=4.2, bottom=0, end=4.2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            border = {"val": "single", "sz": "2", "color": "000000"}
            set_cell_borders(cell, top=border, left=border, bottom=border, right=border)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(13)
            text = labels[cell_index] if row_index == 0 else values[cell_index]
            set_run_font(paragraph.add_run(text), body_font, 12, "808080",
                         bold=row_index == 1)

    add_exact_spacer(18.88)
    for row in payload["rows"]:
        block = document.add_table(rows=1, cols=1)
        set_table_geometry(block, [content_width_pt])
        cell = block.cell(0, 0)
        set_cell_margins(cell, top=14, start=18, bottom=19, end=0)
        set_cell_borders(
            cell,
            top={"val": "nil"},
            left={"val": "single", "sz": "48", "color": row["color"].lstrip("#")},
            bottom={"val": "nil"},
            right={"val": "nil"},
        )
        heading = cell.paragraphs[0]
        heading.paragraph_format.space_after = Pt(7)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        heading.paragraph_format.line_spacing = Pt(20)
        set_run_font(heading.add_run(row["speaker"]), body_font, 16,
                     row["color"].lstrip("#"))
        set_run_font(
            heading.add_run(f"  {row['start_time']} - {row['end_time']}"),
            body_font, 12, "8A8A8A")
        body = cell.add_paragraph()
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        body.paragraph_format.line_spacing = Pt(23)
        set_run_font(body.add_run(row["text"]), body_font, 14, "191919")
        add_exact_spacer(8)

    footer = section.footer
    footer_middle_width = content_width_pt - 112
    footer_table = footer.add_table(rows=1, cols=3, width=Pt(content_width_pt))
    set_table_geometry(footer_table, [100, footer_middle_width, 12])
    footer_table.rows[0].height = Pt(20)
    footer_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    footer_brand_cell, footer_spacer_cell, footer_page_cell = footer_table.rows[0].cells
    for cell in footer_table.rows[0].cells:
        set_cell_margins(cell)
        clear_table_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    footer_brand = footer_brand_cell.paragraphs[0]
    footer_brand.paragraph_format.space_before = Pt(0)
    footer_brand.paragraph_format.space_after = Pt(0)
    set_run_font(footer_brand.add_run("●"), "Arial", 11, "191919")
    set_run_font(footer_brand.add_run("  WordGrab"), brand_font, 14, "8A8A8A")
    footer_spacer = footer_spacer_cell.paragraphs[0]
    footer_spacer.paragraph_format.space_before = Pt(0)
    footer_spacer.paragraph_format.space_after = Pt(0)
    footer_spacer.add_run().add_picture(
        color_strip_image(("#FFFFFF",), width=1000, height=2, gap=0),
        width=Pt(footer_middle_width - 2), height=Pt(1))
    footer_page = footer_page_cell.paragraphs[0]
    footer_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_page.paragraph_format.space_before = Pt(0)
    footer_page.paragraph_format.space_after = Pt(0)
    page_run = footer_page.add_run()
    set_run_font(page_run, body_font, 10, "8A8A8A")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        page_run._r.append(node)
    for paragraph in list(footer.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    document.save(path)


def write_pdf_export(path, payload):
    from html import escape
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer

    font_name = "WordGrabHeiti"
    light_font = "WordGrabHeitiLight"
    brand_font = "WordGrabArialBlack"

    def register_font(name, candidates):
        if name in pdfmetrics.getRegisteredFontNames():
            return name
        for filename, subfont_index in candidates:
            if os.path.isfile(filename):
                try:
                    pdfmetrics.registerFont(
                        TTFont(name, filename, subfontIndex=subfont_index))
                    return name
                except Exception:
                    continue
        return None

    registered_body = register_font(font_name, (
        ("/System/Library/Fonts/STHeiti Medium.ttc", 0),
        ("/System/Library/Fonts/Supplemental/Arial Unicode.ttf", 0),
    ))
    registered_light = register_font(light_font, (
        ("/System/Library/Fonts/STHeiti Light.ttc", 0),
        ("/System/Library/Fonts/STHeiti Medium.ttc", 0),
        ("/System/Library/Fonts/Supplemental/Arial Unicode.ttf", 0),
    ))
    registered_brand = register_font(brand_font, (
        ("/System/Library/Fonts/Supplemental/Arial Black.ttf", 0),
        ("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 0),
    ))
    if not registered_body:
        font_name = "STSong-Light"
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    if not registered_light:
        light_font = font_name
    if not registered_brand:
        brand_font = "Helvetica-Bold"

    body_style = ParagraphStyle(
        "WordGrabBody", fontName=font_name, fontSize=14, leading=23,
        textColor=colors.HexColor("#191919"), wordWrap="CJK",
    )
    heading_style = ParagraphStyle(
        "WordGrabSpeakerHeading", fontName=font_name, fontSize=16, leading=20,
        textColor=colors.HexColor("#191919"), wordWrap="CJK",
    )

    class BrandHeader(Flowable):
        def __init__(self, width):
            super().__init__()
            self.width = width
            self.height = 40

        def wrap(self, available_width, available_height):
            return self.width, self.height

        def draw(self):
            self.canv.drawImage(
                os.path.join(UI, "icon_1024.png"), 0, 0, 40, 40,
                preserveAspectRatio=True, mask="auto")
            self.canv.setFillColor(colors.HexColor("#191919"))
            self.canv.setFont(brand_font, 20)
            self.canv.drawString(47.2, 5.9, "WordGrab")

    class RainbowStrip(Flowable):
        def __init__(self, width):
            super().__init__()
            self.width = width
            self.height = 10

        def wrap(self, available_width, available_height):
            return self.width, self.height

        def draw(self):
            gap = 1
            segment_width = (self.width - gap * (len(EXPORT_COLORS) - 1)) / len(EXPORT_COLORS)
            x = 0
            for color in EXPORT_COLORS:
                self.canv.setFillColor(colors.HexColor(color))
                self.canv.rect(x, 0, segment_width, self.height, stroke=0, fill=1)
                x += segment_width + gap

    class MetadataSummary(Flowable):
        def __init__(self, width):
            super().__init__()
            self.width = width
            self.height = 39.77

        def wrap(self, available_width, available_height):
            return self.width, self.height

        def draw(self):
            canvas = self.canv
            col_width = self.width / 3
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(.2)
            canvas.rect(0, 0, self.width, self.height, stroke=1, fill=0)
            canvas.line(0, 20.58, self.width, 20.58)
            canvas.line(col_width, 0, col_width, self.height)
            canvas.line(col_width * 2, 0, col_width * 2, self.height)
            labels = ("录制时间", "时长 · 字数", "参会人")
            values = (payload["recorded_text"],
                      f"{payload['duration_text']} · {payload['word_count_text']}",
                      payload["participants_text"])
            canvas.setFillColor(colors.HexColor("#808080"))
            canvas.setFont(light_font, 12)
            for index, label in enumerate(labels):
                canvas.drawString(index * col_width + 4.2, 24.5, label)
            canvas.setFont(font_name, 12)
            for index, value in enumerate(values):
                canvas.drawString(index * col_width + 4.2, 5.3, value)

    class SpeakerBlock(Flowable):
        def __init__(self, row, width):
            super().__init__()
            self.row = row
            self.width = width
            self.content_width = width - 18
            heading = (
                f'<font color="{row["color"]}">{escape(row["speaker"])}</font>'
                f'<font size="12" color="#8A8A8A">  '
                f'{escape(row["start_time"])} - {escape(row["end_time"])}</font>'
            )
            self.heading = Paragraph(heading, heading_style)
            self.body = Paragraph(escape(row["text"]), body_style)
            self.heading_height = 0
            self.body_height = 0
            self.height = 0

        def wrap(self, available_width, available_height):
            _, self.heading_height = self.heading.wrap(self.content_width, available_height)
            _, self.body_height = self.body.wrap(self.content_width, available_height)
            self.height = 14 + self.heading_height + 6 + self.body_height + 19
            return self.width, self.height

        def draw(self):
            canvas = self.canv
            canvas.saveState()
            canvas.setStrokeColor(colors.HexColor(self.row["color"]))
            canvas.setLineWidth(6)
            canvas.setLineCap(1)
            canvas.line(0, 3, 0, self.height - 3)
            canvas.restoreState()
            body_y = 19
            self.body.drawOn(canvas, 18, body_y)
            self.heading.drawOn(canvas, 18, body_y + self.body_height + 6)

    page_width, _ = A4
    left_margin, right_margin = 40, 41
    content_width = page_width - left_margin - right_margin
    story = [
        BrandHeader(content_width),
        Spacer(1, 9.68),
        RainbowStrip(content_width),
        Spacer(1, 16.87),
        MetadataSummary(content_width - 1),
        Spacer(1, 18.88),
    ]
    for row in payload["rows"]:
        story.append(SpeakerBlock(row, content_width))
        story.append(Spacer(1, 8))

    document = SimpleDocTemplate(
        path, pagesize=A4, rightMargin=right_margin, leftMargin=left_margin,
        topMargin=14.84, bottomMargin=40,
        title=payload["title"], author="WordGrab",
        subject="WordGrab 本地语音转录稿",
    )

    def draw_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#191919"))
        canvas.circle(left_margin + 6, 16, 5, stroke=0, fill=1)
        canvas.setFillColor(colors.HexColor("#8A8A8A"))
        canvas.setFont(brand_font, 14)
        canvas.drawString(left_margin + 18, 10, "WordGrab")
        canvas.setFont(font_name, 10)
        canvas.drawRightString(A4[0] - right_margin, 18, f"{doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)


def write_ai_pdf_export(path, payload):
    """AI 分析专用报告版式：不复用说话人逐段稿结构。"""
    from html import escape
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (HRFlowable, Image, KeepTogether, Paragraph,
                                   SimpleDocTemplate, Spacer, Table, TableStyle)

    font_name, bold_font = "WordGrabReport", "WordGrabReportBold"

    def register_font(name, candidates):
        if name in pdfmetrics.getRegisteredFontNames(): return name
        for filename, subfont_index in candidates:
            if os.path.isfile(filename):
                try:
                    pdfmetrics.registerFont(TTFont(name, filename, subfontIndex=subfont_index))
                    return name
                except Exception:
                    pass
        return None

    if not register_font(font_name, (("/System/Library/Fonts/STHeiti Light.ttc", 0),
                                     ("/System/Library/Fonts/STHeiti Medium.ttc", 0))):
        font_name = "STSong-Light"; pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    if not register_font(bold_font, (("/System/Library/Fonts/STHeiti Medium.ttc", 0),
                                     ("/System/Library/Fonts/STHeiti Light.ttc", 0))):
        bold_font = font_name

    ink = colors.HexColor("#191919")
    muted = colors.HexColor("#808080")
    accent = colors.HexColor("#12B981")
    secondary = colors.HexColor("#FF9F0A")
    line = colors.HexColor("#D9D9D9")
    soft = colors.white
    blue_soft = colors.HexColor("#F2FBF7")
    green_soft = colors.HexColor("#EFF9F3")
    amber_soft = colors.HexColor("#FFF8E8")

    styles = {
        "eyebrow": ParagraphStyle("AiEyebrow", fontName=bold_font, fontSize=10.5, leading=14,
                                  textColor=accent, spaceAfter=7),
        "title": ParagraphStyle("AiTitle", fontName=bold_font, fontSize=23, leading=31,
                                textColor=ink, spaceAfter=8, wordWrap="CJK"),
        "subtitle": ParagraphStyle("AiSubtitle", fontName=font_name, fontSize=10.5, leading=16,
                                   textColor=muted, spaceAfter=19, wordWrap="CJK"),
        "section": ParagraphStyle("AiSection", fontName=bold_font, fontSize=15, leading=21,
                                  textColor=accent, spaceBefore=17, spaceAfter=9, keepWithNext=True,
                                  wordWrap="CJK"),
        "body": ParagraphStyle("AiBody", fontName=font_name, fontSize=11.2, leading=19,
                               textColor=ink, spaceAfter=8, wordWrap="CJK"),
        "bullet": ParagraphStyle("AiBullet", fontName=font_name, fontSize=11, leading=18,
                                 leftIndent=14, firstLineIndent=-10, textColor=ink,
                                 spaceAfter=6, wordWrap="CJK"),
        "number": ParagraphStyle("AiNumber", fontName=font_name, fontSize=11, leading=18,
                                 leftIndent=23, firstLineIndent=-19, textColor=ink,
                                 spaceAfter=7, wordWrap="CJK"),
        "meta": ParagraphStyle("AiMeta", fontName=font_name, fontSize=9.3, leading=14,
                               textColor=muted, wordWrap="CJK"),
        "quote": ParagraphStyle("AiQuote", fontName=font_name, fontSize=10.8, leading=18,
                                leftIndent=12, rightIndent=8, textColor=ink,
                                spaceAfter=5, wordWrap="CJK"),
        "callout": ParagraphStyle("AiCallout", fontName=bold_font, fontSize=12.2, leading=20,
                                  textColor=ink, wordWrap="CJK"),
        "footer": ParagraphStyle("AiFooter", fontName=font_name, fontSize=8.5, leading=11,
                                 textColor=muted, alignment=TA_CENTER),
    }

    def clean(value):
        text = str(value or "").strip()
        text = re.sub(r"seg-\d+(?:\s*[~～-]\s*seg-\d+)?", "", text, flags=re.I)
        # 综合摘要中的时间证据只用于内部追溯，报告正文不重复展示。
        text = re.sub(r"[（(][0-9:~～\-，,；;、\s]+[）)]", "", text)
        text = re.sub(r"([,，;；]\s*)+[）)]", "）", text)
        text = re.sub(r"[（(]\s*[）)]", "", text)
        text = re.sub(r"\s+([，。；：！？、）])", r"\1", text)
        text = re.sub(r"（\s+", "（", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"[，,；;]\s*[，,；;]", "，", text)
        text = re.sub(r"\s+。", "。", text)
        return text.strip(" ，,;；")

    def item_text(item):
        if isinstance(item, str): return clean(item)
        if not isinstance(item, dict): return ""
        return clean(item.get("text") or item.get("summary") or item.get("quote")
                     or item.get("title") or item.get("task"))

    result = payload.get("ai_result") if isinstance(payload.get("ai_result"), dict) else {}
    overview = result.get("overview") if isinstance(result.get("overview"), dict) else {}
    story = []

    icon_path = os.path.join(UI, "icon_1024.png")
    brand = [[Image(icon_path, 38, 38) if os.path.isfile(icon_path) else "",
              Paragraph("WordGrab", ParagraphStyle("Brand", parent=styles["body"],
                                                     fontName="Helvetica-Bold", fontSize=22, leading=27))]]
    brand_table = Table(brand, colWidths=[45, 180], hAlign="LEFT")
    brand_table.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                                     ("LEFTPADDING", (0,0), (-1,-1), 0),
                                     ("RIGHTPADDING", (0,0), (-1,-1), 0),
                                     ("TOPPADDING", (0,0), (-1,-1), 0),
                                     ("BOTTOMPADDING", (0,0), (-1,-1), 0)]))
    stripe = Table([[""] * len(EXPORT_COLORS)], colWidths=[514 / len(EXPORT_COLORS)] * len(EXPORT_COLORS),
                   rowHeights=[9], hAlign="LEFT")
    stripe_style = [("BACKGROUND", (index, 0), (index, 0), colors.HexColor(color))
                    for index, color in enumerate(EXPORT_COLORS)]
    stripe.setStyle(TableStyle(stripe_style + [("LEFTPADDING", (0,0), (-1,-1), 0),
                                                ("RIGHTPADDING", (0,0), (-1,-1), 0),
                                                ("TOPPADDING", (0,0), (-1,-1), 0),
                                                ("BOTTOMPADDING", (0,0), (-1,-1), 0)]))
    story += [brand_table, Spacer(1, 11), stripe, Spacer(1, 18), Paragraph("AI 内容分析报告", styles["eyebrow"]),
              Paragraph(escape(clean(overview.get("title") or payload.get("analysis_title") or payload.get("source_title"))), styles["title"]),
              Paragraph(escape(clean(payload.get("source_title"))), styles["subtitle"])]

    type_names = {"meeting":"会议", "interview":"访谈", "lecture":"讲座",
                  "call":"通话", "memo":"备忘", "other":"其他"}
    speakers = "、".join(str(x) for x in (overview.get("speakers") or [])) or "—"
    metadata = [
        ["录制时间", payload.get("recorded_text") or "—", "录音时长", overview.get("duration") or payload.get("duration_text") or "—"],
        ["内容类型", type_names.get(overview.get("type"), overview.get("type") or "其他"), "参与人", speakers],
        ["分析模板", payload.get("template_name") or "通用摘要", "分析字数", payload.get("word_count_text") or "—"],
    ]
    meta_data = [[Paragraph(escape(clean(cell)), styles["meta"] if col % 2 == 0 else styles["body"])
                  for col, cell in enumerate(row)] for row in metadata]
    meta_table = Table(meta_data, colWidths=[63, 194, 63, 194], hAlign="LEFT")
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), soft), ("BOX", (0,0), (-1,-1), .55, colors.HexColor("#222222")),
        ("INNERGRID", (0,0), (-1,-1), .45, colors.HexColor("#222222")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 4), ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    story += [meta_table, Spacer(1, 18)]

    one_line = clean(result.get("one_line_summary"))
    if one_line:
        callout = Table([[Paragraph(escape(one_line), styles["callout"])]], colWidths=[514])
        callout.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), blue_soft),
                                     ("LINEBEFORE", (0,0), (0,-1), 5, accent),
                                     ("LEFTPADDING", (0,0), (-1,-1), 14), ("RIGHTPADDING", (0,0), (-1,-1), 14),
                                     ("TOPPADDING", (0,0), (-1,-1), 12), ("BOTTOMPADDING", (0,0), (-1,-1), 12)]))
        story += [Paragraph("核心结论", styles["section"]), callout]

    section_index = 0

    def section_heading(title):
        nonlocal section_index
        module_color = colors.HexColor(EXPORT_COLORS[section_index % len(EXPORT_COLORS)])
        section_index += 1
        heading = Table([[Paragraph(escape(title), styles["section"])]], colWidths=[514])
        heading.setStyle(TableStyle([
            ("LINEBEFORE", (0, 0), (0, -1), 6, module_color),
            ("LEFTPADDING", (0, 0), (-1, -1), 13),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LINEBELOW", (0, 0), (-1, -1), .5, line),
        ]))
        story.extend([Spacer(1, 7), heading, Spacer(1, 9)])

    def add_paragraphs(title, value):
        text = clean(value)
        if not text: return
        section_heading(title)
        parts = [clean(part) for part in re.split(r"\n+", text) if clean(part)]
        story.extend(Paragraph(escape(part), styles["body"]) for part in parts)

    def add_items(title, items, numbered=False, background=None):
        values = [item for item in (items or []) if item_text(item)] if isinstance(items, list) else []
        if not values: return
        section_heading(title)
        for index, item in enumerate(values, 1):
            text = item_text(item)
            when = clean(item.get("time")) if isinstance(item, dict) else ""
            prefix = f"{index}." if numbered else "•"
            suffix = f' <font color="#7A808B">{escape(when)}</font>' if when else ""
            paragraph = Paragraph(f'<font color="#12B981"><b>{prefix}</b></font> {escape(text)}{suffix}',
                                  styles["number"] if numbered else styles["bullet"])
            if background:
                box = Table([[paragraph]], colWidths=[514])
                box.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), background),
                                         ("LEFTPADDING", (0,0), (-1,-1), 11), ("RIGHTPADDING", (0,0), (-1,-1), 11),
                                         ("TOPPADDING", (0,0), (-1,-1), 8), ("BOTTOMPADDING", (0,0), (-1,-1), 4)]))
                story.extend([box, Spacer(1, 5)])
            else:
                story.append(paragraph)

    add_paragraphs("分析摘要", result.get("summary"))
    add_items("内容脉络", result.get("chapters"), numbered=True)
    add_items("关键信息", result.get("key_points"))

    actions = result.get("action_items") if isinstance(result.get("action_items"), list) else []
    if actions:
        section_heading("行动项")
        for index, item in enumerate(actions, 1):
            task = item_text(item); details = []
            if isinstance(item, dict) and item.get("owner"): details.append("负责人：" + clean(item["owner"]))
            due = clean(item.get("due") or item.get("deadline")) if isinstance(item, dict) else ""
            if due and due not in {"未明确", "待确认"}: details.append("截止：" + due)
            body = Paragraph(f'<b>{index}. {escape(task)}</b>' + (f'<br/><font color="#717784">{escape("　".join(details))}</font>' if details else ""), styles["body"])
            box = Table([[body]], colWidths=[514])
            box.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), green_soft),
                                     ("BOX", (0,0), (-1,-1), .45, colors.HexColor("#D6ECDD")),
                                     ("LEFTPADDING", (0,0), (-1,-1), 12), ("RIGHTPADDING", (0,0), (-1,-1), 12),
                                     ("TOPPADDING", (0,0), (-1,-1), 9), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
            story.extend([box, Spacer(1, 6)])

    decisions = result.get("decisions") if isinstance(result.get("decisions"), dict) else {}
    add_items("已确认决策", decisions.get("decided"), background=green_soft)
    add_items("分歧与不同观点", decisions.get("disagreements"), background=amber_soft)
    add_items("待确认问题", decisions.get("open"), background=amber_soft)

    keywords = [clean(x) for x in (result.get("keywords") or []) if clean(x)]
    if keywords:
        section_heading("关键词")
        story.append(Paragraph(escape("　·　".join(keywords)), styles["body"]))

    highlights = result.get("highlights") if isinstance(result.get("highlights"), list) else []
    if highlights:
        section_heading("重要原话")
        for item in highlights:
            quote = item_text(item)
            if not quote: continue
            when = clean(item.get("time")) if isinstance(item, dict) else ""
            story.append(KeepTogether([Paragraph("“" + escape(quote) + "”", styles["quote"]),
                                       Paragraph(escape(when), styles["meta"]) if when else Spacer(1, 0)]))
            story.append(Spacer(1, 7))
    add_items("后续建议", result.get("suggestions"), numbered=True)

    document = SimpleDocTemplate(path, pagesize=A4, leftMargin=40, rightMargin=41,
                                 topMargin=15, bottomMargin=40, title=payload.get("analysis_title") or "AI 内容分析",
                                 author="WordGrab", subject="WordGrab AI 内容分析报告")

    def page_furniture(canvas, doc):
        canvas.saveState()
        if doc.page > 1:
            canvas.setStrokeColor(line); canvas.setLineWidth(.5); canvas.line(40, A4[1] - 27, A4[0] - 41, A4[1] - 27)
            canvas.setFont(font_name, 8.5); canvas.setFillColor(muted)
            canvas.drawString(40, A4[1] - 21, clean(payload.get("analysis_title") or "AI 内容分析"))
        canvas.setFillColor(ink); canvas.circle(46, 16, 5, stroke=0, fill=1)
        canvas.setFont("Helvetica-Bold", 14); canvas.setFillColor(muted)
        canvas.drawString(58, 10, "WordGrab")
        canvas.setFont(font_name, 8.5)
        canvas.drawRightString(A4[0] - 41, 18, str(doc.page))
        canvas.restoreState()

    document.build(story, onFirstPage=page_furniture, onLaterPages=page_furniture)


def write_ai_docx_export(path, payload):
    """AI 分析专用 Word 报告，与品牌 PDF 保持同一视觉系统。"""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin, section.right_margin = Mm(14.1), Mm(14.5)
    section.top_margin, section.bottom_margin = Mm(5.3), Mm(14.1)
    body_font, brand_font = "Arial Unicode MS", "Arial"
    ink, muted, accent, secondary, line = "191919", "808080", "12B981", "FF9F0A", "D9D9D9"

    def clean(value):
        text = str(value or "").strip()
        text = re.sub(r"seg-\d+(?:\s*[~～-]\s*seg-\d+)?", "", text, flags=re.I)
        text = re.sub(r"[（(][0-9:~～\-，,；;、\s]+[）)]", "", text)
        text = re.sub(r"([,，;；]\s*)+[）)]", "）", text)
        text = re.sub(r"[（(]\s*[）)]", "", text)
        text = re.sub(r"\s+([，。；：！？、）])", r"\1", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip(" ，,;；")

    def set_run(run, size=10.5, color=ink, bold=False, name=body_font):
        run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        fonts = run._element.get_or_add_rPr().rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts"); run._element.get_or_add_rPr().insert(0, fonts)
        for key in ("ascii", "hAnsi", "eastAsia"): fonts.set(qn(f"w:{key}"), name)

    def shade(cell, color):
        props = cell._tc.get_or_add_tcPr(); node = props.find(qn("w:shd"))
        if node is None: node = OxmlElement("w:shd"); props.append(node)
        node.set(qn("w:fill"), color)

    def cell_margins(cell, top=90, start=120, bottom=90, end=120):
        props = cell._tc.get_or_add_tcPr(); margins = props.find(qn("w:tcMar"))
        if margins is None: margins = OxmlElement("w:tcMar"); props.append(margins)
        for tag, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
            node = margins.find(qn(f"w:{tag}"))
            if node is None: node = OxmlElement(f"w:{tag}"); margins.append(node)
            node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

    def borders(cell, color=line, size="4", value="single"):
        props = cell._tc.get_or_add_tcPr(); box = props.find(qn("w:tcBorders"))
        if box is None: box = OxmlElement("w:tcBorders"); props.append(box)
        for edge in ("top","left","bottom","right","insideH","insideV"):
            node = box.find(qn(f"w:{edge}"))
            if node is None: node = OxmlElement(f"w:{edge}"); box.append(node)
            node.set(qn("w:val"), value); node.set(qn("w:sz"), size); node.set(qn("w:color"), color)

    def no_borders(table):
        for row in table.rows:
            for cell in row.cells: borders(cell, value="nil")

    def set_width(cell, width_mm):
        width = int(Mm(width_mm).emu / 635)
        tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        if tcw is None: tcw = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcw)
        tcw.set(qn("w:w"), str(width)); tcw.set(qn("w:type"), "dxa")

    def paragraph(text="", size=10.5, color=ink, bold=False, before=0, after=5, leading=1.45):
        p = document.add_paragraph(); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = leading
        set_run(p.add_run(clean(text)), size, color, bold); return p

    def tiny_spacer(points=3):
        p = document.add_paragraph(); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(points)
        set_run(p.add_run("\u200b"), 1, "FFFFFF")

    section_index = 0

    def section_heading(title):
        nonlocal section_index
        module_color = EXPORT_COLORS[section_index % len(EXPORT_COLORS)].lstrip("#")
        section_index += 1
        p = document.add_paragraph(); p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(7); p.paragraph_format.keep_with_next = True
        set_run(p.add_run(title), 14, accent, True)
        p.paragraph_format.left_indent = Mm(4)
        props = p._p.get_or_add_pPr(); box = OxmlElement("w:pBdr")
        left = OxmlElement("w:left"); left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "28")
        left.set(qn("w:color"), module_color); left.set(qn("w:space"), "8"); box.append(left)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4"); bottom.set(qn("w:color"), line)
        bottom.set(qn("w:space"), "6"); box.append(bottom); props.append(box)

    def item_text(item):
        if isinstance(item, str): return clean(item)
        if not isinstance(item, dict): return ""
        return clean(item.get("text") or item.get("summary") or item.get("quote")
                     or item.get("title") or item.get("task"))

    def fresh_numbering_id():
        numbering = document.part.numbering_part.element
        style_num_id = document.styles["List Number"]._element.pPr.numPr.numId.val
        source_num = next(node for node in numbering.findall(qn("w:num"))
                          if node.get(qn("w:numId")) == str(style_num_id))
        abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
        used = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
        new_id = max(used or [0]) + 1
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(new_id))
        abstract = OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"), abstract_id); num.append(abstract)
        override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); num.append(override)
        numbering.append(num); return new_id

    def add_list(title, items, numbered=False, fill=None):
        values = [item for item in (items or []) if item_text(item)] if isinstance(items, list) else []
        if not values: return
        section_heading(title)
        number_id = fresh_numbering_id() if numbered else None
        for index, item in enumerate(values, 1):
            text, when = item_text(item), clean(item.get("time")) if isinstance(item, dict) else ""
            if fill:
                table = document.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.LEFT
                cell = table.cell(0,0); shade(cell, fill); borders(cell, color="D6ECDD" if fill=="EFF9F3" else "F1DFC0")
                cell_margins(cell, 110, 150, 100, 150); p = cell.paragraphs[0]
                set_run(p.add_run("• "), 10.5, accent, True); set_run(p.add_run(text), 10.5)
                if when: set_run(p.add_run("  " + when), 9.2, muted)
                tiny_spacer()
            else:
                p = document.add_paragraph(style="List Number" if numbered else "List Bullet")
                if numbered:
                    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
                    num_pr.get_or_add_ilvl().val = 0; num_pr.get_or_add_numId().val = number_id
                p.paragraph_format.left_indent = Mm(5.5); p.paragraph_format.first_line_indent = Mm(-3.5)
                p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.35
                set_run(p.add_run(text), 10.5)
                if when: set_run(p.add_run("  " + when), 9.2, muted)

    result = payload.get("ai_result") if isinstance(payload.get("ai_result"), dict) else {}
    overview = result.get("overview") if isinstance(result.get("overview"), dict) else {}

    # 品牌标头：Logo + WordGrab + 彩色条，与转写 PDF 保持一致。
    brand = document.add_table(rows=1, cols=2); brand.alignment = WD_TABLE_ALIGNMENT.LEFT; brand.autofit = False
    brand.columns[0].width = Mm(14); brand.columns[1].width = Mm(145)
    brand.cell(0,0).width = Mm(14); brand.cell(0,1).width = Mm(145); no_borders(brand)
    for cell in brand.row_cells(0): cell_margins(cell, 0, 0, 0, 0); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    icon = os.path.join(UI, "icon_1024.png")
    brand.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    brand.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.isfile(icon): brand.cell(0,0).paragraphs[0].add_run().add_picture(icon, width=Mm(12), height=Mm(12))
    set_run(brand.cell(0,1).paragraphs[0].add_run("WordGrab"), 22, ink, True, brand_font)

    stripe = document.add_table(rows=1, cols=len(EXPORT_COLORS)); stripe.alignment = WD_TABLE_ALIGNMENT.LEFT
    stripe.rows[0].height = Pt(8); stripe.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell, color in zip(stripe.rows[0].cells, EXPORT_COLORS):
        shade(cell, color.lstrip("#")); borders(cell, color="FFFFFF", size="6"); cell_margins(cell,0,0,0,0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    p = document.add_paragraph(); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run("AI 内容分析报告"), 9.5, accent, True)
    p = document.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(clean(overview.get("title") or payload.get("analysis_title") or payload.get("source_title"))), 23, ink, True)
    p = document.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    set_run(p.add_run(clean(payload.get("source_title"))), 10, muted)

    type_names = {"meeting":"会议","interview":"访谈","lecture":"讲座","call":"通话","memo":"备忘","other":"其他"}
    speakers = "、".join(str(x) for x in (overview.get("speakers") or [])) or "—"
    meta = [
        ["录制时间", payload.get("recorded_text") or "—", "录音时长", overview.get("duration") or payload.get("duration_text") or "—"],
        ["内容类型", type_names.get(overview.get("type"), overview.get("type") or "其他"), "参与人", speakers],
        ["分析模板", payload.get("template_name") or "通用摘要", "分析字数", payload.get("word_count_text") or "—"],
    ]
    table = document.add_table(rows=3, cols=4); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [24, 62, 24, 62]
    for r, values in enumerate(meta):
        for c, value in enumerate(values):
            cell=table.cell(r,c); set_width(cell,widths[c]); shade(cell,"FFFFFF"); borders(cell,color="222222",size="4")
            cell_margins(cell,55,65,55,65); cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            p=cell.paragraphs[0]; set_run(p.add_run(clean(value)),9 if c%2==0 else 10,muted if c%2==0 else ink,False)

    one_line = clean(result.get("one_line_summary"))
    if one_line:
        section_heading("核心结论")
        box=document.add_table(rows=1,cols=1); cell=box.cell(0,0); shade(cell,"F2FBF7"); borders(cell,color=accent)
        cell_margins(cell,150,180,150,180); set_run(cell.paragraphs[0].add_run(one_line),12,ink,True)

    summary=clean(result.get("summary"))
    if summary:
        section_heading("分析摘要")
        for part in [clean(x) for x in re.split(r"\n+",summary) if clean(x)]: paragraph(part)
    add_list("内容脉络",result.get("chapters"),numbered=True)
    add_list("关键信息",result.get("key_points"))

    actions=result.get("action_items") if isinstance(result.get("action_items"),list) else []
    if actions:
        section_heading("行动项")
        for index,item in enumerate(actions,1):
            table=document.add_table(rows=1,cols=1); cell=table.cell(0,0); shade(cell,"EFF9F3"); borders(cell,color="D6ECDD")
            cell_margins(cell,120,160,110,160); p=cell.paragraphs[0]
            set_run(p.add_run(f"{index}. {item_text(item)}"),10.5,ink,True)
            details=[]
            if isinstance(item,dict) and item.get("owner"): details.append("负责人："+clean(item["owner"]))
            due=clean(item.get("due") or item.get("deadline")) if isinstance(item,dict) else ""
            if due and due not in {"未明确","待确认"}: details.append("截止："+due)
            if details: p.add_run("\n"); set_run(p.add_run("　".join(details)),9.2,muted)
            tiny_spacer()

    decisions=result.get("decisions") if isinstance(result.get("decisions"),dict) else {}
    add_list("已确认决策",decisions.get("decided"),fill="EFF9F3")
    add_list("分歧与不同观点",decisions.get("disagreements"),fill="FFF8E8")
    add_list("待确认问题",decisions.get("open"),fill="FFF8E8")

    keywords=[clean(x) for x in (result.get("keywords") or []) if clean(x)]
    if keywords: section_heading("关键词"); paragraph("　·　".join(keywords))
    highlights=result.get("highlights") if isinstance(result.get("highlights"),list) else []
    if highlights:
        section_heading("重要原话")
        for item in highlights:
            quote=item_text(item)
            if not quote: continue
            p=paragraph("“"+quote+"”",10.5,ink,False,0,4)
            p.paragraph_format.left_indent=Mm(5)
            when=clean(item.get("time")) if isinstance(item,dict) else ""
            if when: set_run(p.add_run("  " + when),9,muted)
    add_list("后续建议",result.get("suggestions"),numbered=True)

    footer=section.footer
    footer.paragraphs[0]._element.getparent().remove(footer.paragraphs[0]._element)
    footer_table=footer.add_table(rows=1,cols=2,width=Mm(181)); footer_table.autofit=False
    footer_table.columns[0].width=Mm(150); footer_table.columns[1].width=Mm(31)
    footer_table.cell(0,0).width=Mm(150); footer_table.cell(0,1).width=Mm(31); no_borders(footer_table)
    for cell in footer_table.rows[0].cells: cell_margins(cell,0,0,0,0)
    left=footer_table.cell(0,0).paragraphs[0]
    set_run(left.add_run("●  WordGrab"),12,muted,True,brand_font)
    right=footer_table.cell(0,1).paragraphs[0]; right.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); right._p.append(fld)

    document.core_properties.title=clean(payload.get("analysis_title") or "AI 内容分析")
    document.core_properties.subject="WordGrab AI 内容分析报告"
    document.core_properties.author="WordGrab"
    document.save(path)


def item_dir(iid):
    return os.path.join(DATA, iid)


def record_export_path(iid, resource, path):
    """记录最近一次用户导出的成品文件，不暴露内部 JSON 存储。"""
    if resource not in {"document", "summary"} or not path:
        return
    record_path = os.path.join(item_dir(iid), "exports.json")
    try:
        with open(record_path, encoding="utf-8") as file:
            records = json.load(file)
        if not isinstance(records, dict):
            records = {}
    except (OSError, ValueError, TypeError):
        records = {}
    records[resource] = {
        "path": os.path.abspath(os.path.expanduser(str(path))),
        "updated_at": datetime.datetime.now().isoformat(timespec="seconds"),
    }
    atomic_write_json(record_path, records)


def latest_export_path(iid, resource):
    if resource not in {"document", "summary"}:
        return None
    try:
        with open(os.path.join(item_dir(iid), "exports.json"), encoding="utf-8") as file:
            value = json.load(file).get(resource)
        return value.get("path") if isinstance(value, dict) else None
    except (OSError, ValueError, TypeError, AttributeError):
        return None


def load_item(iid):
    with open(os.path.join(item_dir(iid), "transcript.json"), encoding="utf-8") as f:
        return json.load(f)


def save_item(iid, data):
    path = os.path.join(item_dir(iid), "transcript.json")
    temporary = path + ".tmp"
    with open(temporary, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.flush()
        os.fsync(f.fileno())
    os.replace(temporary, path)


def _valid_hex_color(value):
    if not isinstance(value, str):
        return None
    color = value.strip().upper()
    if re.fullmatch(r"#[0-9A-F]{6}", color):
        return color
    return None


def _speaker_indexes(data):
    indexes = set()
    for segment in data.get("segments", []) or []:
        try:
            indexes.add(int(segment.get("spk", 0) or 0))
        except (TypeError, ValueError):
            indexes.add(0)
    for key in (data.get("speakers") or {}).keys():
        try:
            indexes.add(int(key))
        except (TypeError, ValueError):
            pass
    return sorted(indexes) or [0]


def make_speaker_colors(seed, indexes):
    seed = str(seed or uuid.uuid4().hex)
    ranked = sorted(
        (hashlib.sha256(f"{seed}:{color}".encode("utf-8")).hexdigest(), color)
        for color in SPEAKER_COLOR_POOL
    )
    palette = [color for _, color in ranked]
    return {
        str(index): palette[position % len(palette)]
        for position, index in enumerate(sorted(int(i) for i in indexes))
    }


def ensure_speaker_colors(iid, data):
    raw = data.get("speaker_colors")
    existing = raw if isinstance(raw, dict) else {}
    changed = not isinstance(raw, dict)
    colors = {}
    generated = make_speaker_colors(iid, _speaker_indexes(data))
    for key, generated_color in generated.items():
        color = _valid_hex_color(existing.get(key))
        colors[key] = color or generated_color
        changed = changed or color is None
    for key, value in existing.items():
        if key not in colors:
            color = _valid_hex_color(value)
            if color:
                colors[key] = color
    if data.get("speaker_colors") != colors:
        data["speaker_colors"] = colors
        changed = True
    return colors, changed


def audio_path_of(iid):
    d = item_dir(iid)
    for fn in os.listdir(d):
        if fn.startswith("audio"):
            return os.path.join(d, fn)
    return None


# ---------- 本地 HTTP 服务（UI + 支持 Range 的音频） ----------

class Handler(BaseHTTPRequestHandler):
    def log_message(self, *a):
        pass

    def _send_json(self, value, status=200):
        body = json.dumps(value, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_jsonp(self, callback, value):
        if not re.fullmatch(r"[A-Za-z_$][\w$\.]{0,120}", callback or ""):
            return self.send_error(400, "无效的回调名称")
        body = f"{callback}({json.dumps(value, ensure_ascii=False)});".encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/javascript; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _read_json(self):
        try:
            length = int(self.headers.get("Content-Length", 0) or 0)
        except ValueError:
            raise ValueError("无效的请求大小")
        if length < 0 or length > 65536:
            raise ValueError("请求内容过大")
        if not length:
            return {}
        value = json.loads(self.rfile.read(length).decode("utf-8"))
        return value if isinstance(value, dict) else {}

    def _send_file(self, path, ctype):
        try:
            size = os.path.getsize(path)
        except OSError:
            self.send_error(404); return
        rng = self.headers.get("Range")
        start, end = 0, size - 1
        if rng:
            m = re.match(r"bytes=(\d+)-(\d*)", rng)
            if m:
                start = int(m.group(1))
                if m.group(2):
                    end = int(m.group(2))
        length = end - start + 1
        self.send_response(206 if rng else 200)
        self.send_header("Content-Type", ctype)
        self.send_header("Accept-Ranges", "bytes")
        self.send_header("Content-Length", str(length))
        if rng:
            self.send_header("Content-Range", f"bytes {start}-{end}/{size}")
        self.end_headers()
        with open(path, "rb") as f:
            f.seek(start)
            remaining = length
            while remaining > 0:
                chunk = f.read(min(65536, remaining))
                if not chunk:
                    break
                try:
                    self.wfile.write(chunk)
                except (BrokenPipeError, ConnectionResetError):
                    break
                remaining -= len(chunk)

    def do_GET(self):
        from urllib.parse import urlparse, parse_qs
        parsed = urlparse(self.path)
        path = parsed.path
        query = parse_qs(parsed.query)
        if path == "/" or path == "/index.html":
            return self._send_file(os.path.join(UI, "index.html"), "text/html; charset=utf-8")
        if path.startswith("/static/"):
            fn = os.path.basename(path)
            fp = os.path.join(UI, fn)
            if os.path.exists(fp):
                ext = os.path.splitext(fn)[1].lower()
                ctype = {
                    ".css": "text/css; charset=utf-8",
                    ".js": "application/javascript; charset=utf-8",
                    ".png": "image/png",
                    ".svg": "image/svg+xml",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".webp": "image/webp",
                }.get(ext, "application/octet-stream")
                return self._send_file(fp, ctype)
        if path.startswith("/audio/"):
            iid = path.split("/audio/")[1]
            ap = audio_path_of(iid)
            if ap:
                ext = os.path.splitext(ap)[1].lower()
                return self._send_file(ap, AUDIO_EXT.get(ext, "application/octet-stream"))
        if path.startswith("/status/"):
            iid = path.split("/status/")[1]
            st = PROGRESS.get(iid) or {"status": "unknown"}
            body = json.dumps(st, ensure_ascii=False).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        if path.startswith("/ai/task/") and API_REF is not None:
            task_id = path.split("/ai/task/", 1)[1]
            result = API_REF.get_ai_task(task_id)
            callback = query.get("callback", [""])[0]
            return self._send_jsonp(callback, result) if callback else self._send_json(result)
        if path == "/ai/start" and API_REF is not None:
            result = API_REF.start_ai_summary(
                query.get("item_id", [""])[0], query.get("template", ["general"])[0],
                query.get("privacy_confirmed", ["0"])[0] == "1",
            )
            callback = query.get("callback", [""])[0]
            return self._send_jsonp(callback, result) if callback else self._send_json(result)
        if path == "/ai/cancel" and API_REF is not None:
            result = API_REF.cancel_ai_task(query.get("task_id", [""])[0])
            callback = query.get("callback", [""])[0]
            return self._send_jsonp(callback, result) if callback else self._send_json(result)
        self.send_error(404)

    def do_POST(self):
        # 拖拽上传：WebView 里拿不到文件路径，只能把字节流 POST 过来落盘再转写
        from urllib.parse import urlparse, parse_qs, unquote
        parsed = urlparse(self.path)
        if parsed.path == "/ai/start" and API_REF is not None:
            try:
                values = self._read_json()
                result = API_REF.start_ai_summary(
                    values.get("item_id"), values.get("template", "general"),
                    bool(values.get("privacy_confirmed", False)),
                )
                return self._send_json(result)
            except (ValueError, json.JSONDecodeError) as exc:
                return self._send_json({"ok": False, "message": str(exc)}, 400)
        if parsed.path == "/ai/cancel" and API_REF is not None:
            try:
                values = self._read_json()
                return self._send_json(API_REF.cancel_ai_task(values.get("task_id")))
            except (ValueError, json.JSONDecodeError) as exc:
                return self._send_json({"ok": False, "message": str(exc)}, 400)
        if parsed.path != "/upload" or API_REF is None:
            return self.send_error(404)
        q = parse_qs(parsed.query)
        name = unquote(q.get("name", ["audio"])[0])
        try:
            length = int(self.headers.get("Content-Length", 0) or 0)
        except ValueError:
            return self.send_error(400, "无效的文件大小")
        if length <= 0 or length > MAX_UPLOAD_BYTES:
            return self.send_error(413, "文件过大，单个文件不能超过 2GB")
        iid = uuid.uuid4().hex[:12]
        os.makedirs(item_dir(iid), exist_ok=True)
        ext = os.path.splitext(name)[1].lower() or ".m4a"
        dst = os.path.join(item_dir(iid), "audio" + ext)
        try:
            remaining = length
            with open(dst, "wb") as f:
                while remaining > 0:
                    chunk = self.rfile.read(min(1 << 20, remaining))
                    if not chunk:
                        break
                    f.write(chunk)
                    remaining -= len(chunk)
            if remaining:
                raise OSError("上传未完成")
        except Exception as e:
            shutil.rmtree(item_dir(iid), ignore_errors=True)
            self.send_error(500, str(e)); return
        title = os.path.splitext(os.path.basename(name))[0]
        PROGRESS[iid] = {"stage": "准备中…", "pct": None, "info": None, "status": "running",
                         "msg": "", "title": title, "partial": []}
        threading.Thread(target=API_REF._run_transcribe, args=(iid, dst, title),
                         daemon=True).start()
        body = json.dumps({"id": iid, "title": title}).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def start_server():
    srv = ThreadingHTTPServer(("127.0.0.1", 0), Handler)
    port = srv.server_address[1]
    threading.Thread(target=srv.serve_forever, daemon=True).start()
    return port


# ---------- JS ↔ Python 桥 ----------

class Api:
    def __init__(self):
        self.window = None
        self.port = None

    # 历史列表
    def list_items(self):
        items = load_index()
        for item in items:
            progress = PROGRESS.get(item.get("id"))
            if progress and progress.get("status") in {"running", "draft", "error"}:
                item["status"] = progress.get("status")
                if progress.get("status") == "error":
                    item["error"] = progress.get("msg", "")
        return items

    # 前端埋点日志
    def log(self, msg):
        print(f"[JS] {msg}", flush=True)
        return True

    # 设置（写入本地 data/settings.json，重启后继续生效）
    def get_settings(self):
        return load_settings()

    def update_settings(self, patch):
        return apply_settings_patch(patch)

    # AI 总结：只支持用户配置的 OpenAI 兼容接口，Key 存在本机受限文件中。
    def get_ai_settings(self):
        settings = load_settings()
        key = ai_service.load_api_key()
        return {
            "base_url": settings["ai_base_url"],
            "model": settings["ai_model"],
            "summary_template": settings["ai_summary_template"],
            "key_configured": bool(key),
            "key_last4": key[-4:] if key else "",
        }

    def list_ai_templates(self):
        path = os.path.join(DATA, "ai_templates.json")
        try:
            with open(path, encoding="utf-8") as file:
                custom = json.load(file)
        except (OSError, ValueError):
            custom = []
        if not isinstance(custom, list):
            custom = []
        return [{"id": "general", "name": "通用摘要", "builtin": True,
                 "objective": "自动识别录音内容，完整整理主题、结论、决策、行动与重要信息",
                 "focus": ["内容概述", "关键结论", "重要信息", "决策与共识", "行动项", "风险与待确认问题", "后续建议", "关键词"],
                 "instructions": "所有事实必须来自原文，并尽量关联时间与原文位置", "detail": "standard"}] + custom

    def save_ai_template(self, values):
        values = values if isinstance(values, dict) else {}
        name = str(values.get("name") or "").strip()[:30]
        objective = str(values.get("objective") or "").strip()[:2000]
        if not name or not objective:
            return {"ok": False, "message": "请填写模板名称和分析目标"}
        templates = [item for item in self.list_ai_templates() if not item.get("builtin")]
        template_id = str(values.get("id") or "").strip()
        if not template_id or template_id == "general":
            template_id = "custom-" + uuid.uuid4().hex[:12]
        if any(item.get("name") == name and item.get("id") != template_id for item in templates):
            return {"ok": False, "message": "已经存在同名模板"}
        focus = values.get("focus") if isinstance(values.get("focus"), list) else []
        payload = {"id": template_id, "name": name, "builtin": False,
                   "objective": objective, "focus": [str(x).strip()[:60] for x in focus if str(x).strip()][:20],
                   "instructions": str(values.get("instructions") or "").strip()[:3000],
                   "detail": values.get("detail") if values.get("detail") in {"concise", "standard", "detailed"} else "standard",
                   "updated_at": datetime.datetime.now().isoformat(timespec="seconds")}
        replaced = False
        for index, item in enumerate(templates):
            if item.get("id") == template_id:
                templates[index], replaced = payload, True
                break
        if not replaced:
            templates.append(payload)
        atomic_write_json(os.path.join(DATA, "ai_templates.json"), templates)
        return {"ok": True, "template": payload, "templates": self.list_ai_templates()}

    def delete_ai_template(self, template_id):
        if not template_id or template_id == "general":
            return {"ok": False, "message": "通用摘要不能删除"}
        templates = [item for item in self.list_ai_templates()
                     if not item.get("builtin") and item.get("id") != template_id]
        atomic_write_json(os.path.join(DATA, "ai_templates.json"), templates)
        if load_settings().get("ai_summary_template") == template_id:
            apply_settings_patch({"ai_summary_template": "general"})
        return {"ok": True, "templates": self.list_ai_templates()}

    def save_ai_settings(self, values):
        values = values if isinstance(values, dict) else {}
        try:
            base_url = ai_service.normalize_base_url(values.get("base_url"))
            model = str(values.get("model") or "").strip()
            if not model:
                return {"ok": False, "code": "AI_MODEL_MISSING", "message": "请输入模型名称"}
            key = str(values.get("api_key") or "").strip()
            last4 = ai_service.save_api_key(key) if key else (ai_service.load_api_key()[-4:] or "")
            if not last4:
                return {"ok": False, "code": "AI_KEY_MISSING", "message": "请输入 API Key"}
            template = str(values.get("summary_template") or "general")
            if not any(item.get("id") == template for item in self.list_ai_templates()):
                template = "general"
            saved = apply_settings_patch({
                "ai_base_url": base_url,
                "ai_model": model,
                "ai_summary_template": template,
            })
            return {"ok": True, "base_url": saved["ai_base_url"], "model": saved["ai_model"],
                    "summary_template": saved["ai_summary_template"], "key_last4": last4}
        except ai_service.AiServiceError as exc:
            return {"ok": False, "code": exc.code, "message": exc.message}
        except OSError:
            return {"ok": False, "code": "AI_KEY_STORAGE_ERROR", "message": "无法在本机安全保存 API Key"}

    def delete_ai_key(self):
        ai_service.delete_api_key()
        return {"ok": True}

    def list_ai_models(self, base_url=None, api_key=None):
        settings = load_settings()
        url = str(base_url or settings["ai_base_url"] or "")
        key = str(api_key or "").strip() or ai_service.load_api_key()
        if not key:
            return {"ok": False, "code": "AI_KEY_MISSING", "message": "请输入 API Key"}
        try:
            return {"ok": True, "models": ai_service.list_models(url, key)}
        except ai_service.AiServiceError as exc:
            return {"ok": False, "code": exc.code, "message": exc.message}

    def test_ai_connection(self, values):
        values = values if isinstance(values, dict) else {}
        settings = load_settings()
        url = str(values.get("base_url") or settings["ai_base_url"] or "")
        model = str(values.get("model") or settings["ai_model"] or "").strip()
        key = str(values.get("api_key") or "").strip() or ai_service.load_api_key()
        if not key:
            return {"ok": False, "code": "AI_KEY_MISSING", "message": "请输入 API Key"}
        if not model:
            return {"ok": False, "code": "AI_MODEL_MISSING", "message": "请输入模型名称"}
        started = time.monotonic()
        try:
            ai_service.test_connection(url, key, model)
            return {"ok": True, "elapsed": round(time.monotonic() - started, 1), "model": model}
        except ai_service.AiServiceError as exc:
            return {"ok": False, "code": exc.code, "message": exc.message}

    def get_ai_summary(self, iid):
        if not iid or not os.path.isdir(item_dir(iid)):
            return None
        return ai_service.load_summary(item_dir(iid))

    def save_ai_summary(self, iid, result):
        existing = self.get_ai_summary(iid)
        if not existing or not isinstance(result, dict):
            return {"ok": False, "message": "没有可保存的总结"}
        existing["result"] = result
        existing["edited"] = True
        existing["updated_at"] = datetime.datetime.now().isoformat(timespec="seconds")
        ai_service.save_summary(item_dir(iid), existing)
        return {"ok": True, "summary": existing}

    def start_ai_summary(self, iid, template="general", privacy_confirmed=False):
        if not iid or not os.path.isfile(os.path.join(item_dir(iid), "transcript.json")):
            return {"ok": False, "code": "AI_TRANSCRIPT_MISSING", "message": "当前文稿尚未保存完成"}
        settings = load_settings()
        base_url, model = settings["ai_base_url"], settings["ai_model"]
        key = ai_service.load_api_key()
        if not base_url or not model or not key:
            return {"ok": False, "code": "AI_NOT_CONFIGURED", "message": "请先在设置中配置 AI 服务"}
        try:
            from urllib.parse import urlparse
            host = urlparse(ai_service.normalize_base_url(base_url)).netloc
        except ai_service.AiServiceError as exc:
            return {"ok": False, "code": exc.code, "message": exc.message}
        if settings.get("ai_privacy_host") != host and not privacy_confirmed:
            return {"ok": False, "code": "AI_PRIVACY_CONFIRM_REQUIRED", "host": host,
                    "message": "生成总结会把当前文稿文字发送到该 AI 服务"}
        if privacy_confirmed:
            apply_settings_patch({"ai_privacy_host": host})
        template_config = next((item for item in self.list_ai_templates() if item.get("id") == template), None)
        if not template_config:
            template, template_config = "general", self.list_ai_templates()[0]
        task_id = uuid.uuid4().hex[:16]
        with AI_TASKS_LOCK:
            AI_TASKS[task_id] = {"id": task_id, "item_id": iid, "status": "queued",
                                 "stage": "正在准备文稿", "current": 0, "total": 1,
                                 "cancelled": False, "message": "", "started_at": time.time()}
        threading.Thread(target=self._run_ai_summary,
                         args=(task_id, iid, template, template_config, base_url, key, model), daemon=True).start()
        return {"ok": True, "task_id": task_id}

    def _run_ai_summary(self, task_id, iid, template, template_config, base_url, key, model):
        def cancelled():
            with AI_TASKS_LOCK:
                return bool(AI_TASKS.get(task_id, {}).get("cancelled"))

        def progress(current, total, stage):
            with AI_TASKS_LOCK:
                task = AI_TASKS.get(task_id)
                if task:
                    task.update(status="running", current=current, total=max(1, total), stage=stage)
        try:
            data = load_item(iid)
            transcript_hash = hashlib.sha256(
                json.dumps(data.get("segments") or [], ensure_ascii=False, sort_keys=True).encode("utf-8")
            ).hexdigest()
            result = ai_service.generate_summary(data, template_config, base_url, key, model,
                                                 progress=progress, cancelled=cancelled)
            quality_warning = result.pop("_quality_warning", "")
            payload = {
                "template": template,
                "template_name": template_config.get("name") or "通用摘要",
                "template_snapshot": template_config,
                "model": model,
                "created_at": datetime.datetime.now().isoformat(timespec="seconds"),
                "updated_at": datetime.datetime.now().isoformat(timespec="seconds"),
                "transcript_hash": transcript_hash,
                "edited": False,
                "quality_warning": quality_warning,
                "result": result,
            }
            ai_service.save_summary(item_dir(iid), payload)
            with AI_TASKS_LOCK:
                AI_TASKS[task_id].update(status="done", current=1, total=1,
                                         stage="总结已完成", summary=payload,
                                         elapsed_seconds=round(time.time() - AI_TASKS[task_id]["started_at"], 1))
            print(f"[ai] 总结完成 task={task_id} elapsed={time.time() - AI_TASKS[task_id]['started_at']:.1f}s", flush=True)
        except ai_service.AiServiceError as exc:
            with AI_TASKS_LOCK:
                AI_TASKS[task_id].update(status="cancelled" if exc.code == "AI_CANCELLED" else "error",
                                         code=exc.code, message=exc.message, stage=exc.message)
            print(f"[ai] 总结失败 task={task_id} code={exc.code} message={exc.message}", flush=True)
        except Exception as exc:
            print(f"[ai] 总结失败: {exc!r}", flush=True)
            with AI_TASKS_LOCK:
                AI_TASKS[task_id].update(status="error", code="AI_INTERNAL_ERROR",
                                         message="生成总结时发生内部错误", stage="生成失败")

    def get_ai_task(self, task_id):
        with AI_TASKS_LOCK:
            task = AI_TASKS.get(task_id)
            if not task:
                return {"status": "unknown"}
            result = dict(task)
            result["elapsed_seconds"] = round(time.time() - task.get("started_at", time.time()), 1)
            return result

    def cancel_ai_task(self, task_id):
        with AI_TASKS_LOCK:
            if task_id in AI_TASKS:
                AI_TASKS[task_id]["cancelled"] = True
                AI_TASKS[task_id]["stage"] = "正在取消"
                return {"ok": True}
        return {"ok": False, "message": "任务不存在"}

    def set_theme(self, theme):
        if theme not in THEME_KEYS:
            return False
        apply_settings_patch({"theme": theme})
        return True

    def pick_export_directory(self):
        import webview
        current = load_settings()["export_directory"]
        try:
            res = self.window.create_file_dialog(webview.FOLDER_DIALOG, directory=current)
        except Exception as e:
            print(f"[bridge] 选择导出目录失败: {e!r}", flush=True)
            return None
        if not res:
            return None
        directory = res[0] if isinstance(res, (list, tuple)) else res
        if directory and os.path.isdir(directory):
            apply_settings_patch({"export_directory": directory})
            return directory
        return None

    def pick_data_directory(self):
        """只选择目录；实际搬迁要等前端二次确认后再执行。"""
        import webview
        current_parent = os.path.dirname(DATA)
        start_directory = current_parent if os.path.isdir(current_parent) else DATA
        try:
            res = self.window.create_file_dialog(
                webview.FOLDER_DIALOG,
                directory=start_directory,
            )
        except Exception as e:
            print(f"[bridge] 选择数据目录失败: {e!r}", flush=True)
            return None
        if not res:
            return None
        directory = res[0] if isinstance(res, (list, tuple)) else res
        if directory and os.path.isdir(directory):
            return os.path.abspath(os.path.expanduser(directory))
        return None

    def set_data_directory(self, directory):
        return migrate_data_directory(directory)

    def get_system_info(self):
        import engine
        try:
            ffmpeg_path = engine._resolve_ffmpeg()
            ffmpeg_ok = True
        except Exception:
            ffmpeg_path = "未找到"
            ffmpeg_ok = False
        caches = model_cache_dirs()
        return {
            "version": APP_VERSION,
            "data_path": DATA,
            "data_size": path_size(DATA),
            "model_path": os.path.expanduser("~/.cache/modelscope"),
            "model_size": sum(path_size(path) for path in caches),
            "model_ready": bool(caches),
            "ffmpeg_ok": ffmpeg_ok,
            "ffmpeg_path": ffmpeg_path,
        }

    def open_local_resource(self, resource):
        allowed = {
            "data": DATA,
            "models": os.path.expanduser("~/.cache/modelscope"),
            "readme": os.path.join(HERE, "README.md"),
            "license": os.path.join(HERE, "LICENSE"),
        }
        path = allowed.get(resource)
        if not path or not os.path.exists(path):
            return False
        try:
            subprocess.Popen(["open", path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return True
        except Exception:
            return False

    def reveal_item_resource(self, iid, resource):
        """在 Finder 中定位最近一次导出的成品，而不是内部 JSON。"""
        resource = str(resource)
        if resource not in {"document", "summary"} or not iid:
            return {"ok": False, "message": "无法识别要显示的文件"}
        path = latest_export_path(str(iid), resource)
        label = "文稿" if resource == "document" else "AI 总结"
        if not path:
            return {"ok": False, "code": "EXPORT_NOT_FOUND", "needs_export": True,
                    "message": f"尚未导出{label}，请先选择保存位置"}
        if not os.path.isfile(path):
            return {"ok": False, "code": "EXPORT_MOVED", "needs_export": True,
                    "message": f"上次导出的{label}已被移动或删除，请重新导出"}
        try:
            subprocess.Popen(["open", "-R", path], stdout=subprocess.DEVNULL,
                             stderr=subprocess.DEVNULL)
            print(f"[finder] 显示导出{label}: {path}", flush=True)
            return {"ok": True, "path": path, "filename": os.path.basename(path)}
        except Exception:
            return {"ok": False, "message": "无法打开 Finder"}

    # 打开某条：返回合并后的分段 + 说话人名 + 音频地址
    def open_item(self, iid):
        import engine
        data = load_item(iid)
        speaker_colors, colors_changed = ensure_speaker_colors(iid, data)
        if colors_changed:
            save_item(iid, data)
        spk_pending = bool(data.get("spk_pending"))
        # 初稿全部 spk=0，按说话人合并会揉成一大段，保留分块粒度
        merged = data["segments"] if spk_pending else engine.merge_by_speaker(data["segments"])
        meta = next((x for x in load_index() if x["id"] == iid), {})
        audio_path = audio_path_of(iid)
        audio_format = (os.path.splitext(audio_path)[1].lstrip(".").upper()
                        if audio_path else "音频")
        return {
            "id": iid,
            "title": meta.get("title", ""),
            "duration": meta.get("duration", 0),
            "created": meta.get("created", ""),
            "audio_format": audio_format,
            "speakers": data.get("speakers", {}),
            "speaker_colors": speaker_colors,
            "segments": merged,
            "spk_pending": spk_pending,
            "audio_url": f"http://127.0.0.1:{self.port}/audio/{iid}",
        }

    # 选文件（原生对话框）
    def pick_file(self):
        import webview
        print("[bridge] pick_file 进入", flush=True)
        try:
            res = self.window.create_file_dialog(webview.OPEN_DIALOG, allow_multiple=False)
            print(f"[bridge] pick_file 返回: {res}", flush=True)
        except Exception as e:
            print(f"[bridge] pick_file 异常: {e}", flush=True)
            return None
        if res:
            return res[0] if isinstance(res, (list, tuple)) else res
        return None

    # 开始转写（后台线程；进度写入 PROGRESS，前端轮询 /status/<iid>）
    def start_transcribe(self, src_path):
        iid = uuid.uuid4().hex[:12]
        title = os.path.splitext(os.path.basename(src_path))[0]
        with INDEX_LOCK:
            items = load_index()
            items.insert(0, {"id": iid, "title": title, "duration": 0,
                             "created": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                             "n_speakers": 0, "status": "queued"})
            save_index(items)
        PROGRESS[iid] = {"stage": "准备中…", "pct": None, "info": None, "status": "running",
                        "msg": "", "title": title, "partial": []}
        threading.Thread(target=self._do_transcribe, args=(iid, src_path), daemon=True).start()
        return iid

    @staticmethod
    def _set(iid, stage, pct=None, info=None, status=None, msg=None, **extra):
        # 合并更新（partial/title 等字段跨多次 _set 保留）
        st = PROGRESS.setdefault(iid, {})
        st.update({"stage": stage, "pct": pct})
        if info is not None:
            st["info"] = {**(st.get("info") or {}), **info}
        if status is not None:
            st["status"] = status
        if msg is not None:
            st["msg"] = msg
        st.update(extra)

    def _do_transcribe(self, iid, src_path):
        # 走「导入按钮」：先把源文件复制进条目目录，再复用转写流程
        try:
            os.makedirs(item_dir(iid), exist_ok=True)
            ext = os.path.splitext(src_path)[1].lower() or ".m4a"
            dst = os.path.join(item_dir(iid), "audio" + ext)
            shutil.copy2(src_path, dst)
            title = os.path.splitext(os.path.basename(src_path))[0]
        except Exception as e:
            self._set(iid, "出错", status="error", msg=str(e))
            return
        self._run_transcribe(iid, dst, title)

    def _run_transcribe(self, iid, audio_file, title):
        # 对已落盘的 audio_file 执行转写（导入按钮 / 拖拽上传共用）
        # 两阶段：①快路径初稿（渐进出字，无说话人）→ 存盘可读；②声纹分离 → 更新说话人
        import engine
        settings = load_settings()
        auto_diarization = settings["auto_diarization"]
        transcription_mode = settings["transcription_mode"]
        print(f"[transcribe] start iid={iid} file={audio_file}", flush=True)
        if TRANSCRIBE_LOCK.locked():
            self._set(iid, "排队中，等待上一个转写完成…")
        with TRANSCRIBE_LOCK:
            wav = None
            try:
                def prog(stage, pct, info=None):
                    print(f"[transcribe] {iid} · {stage} pct={pct}", flush=True)
                    self._set(iid, stage, pct, info)

                dur = engine.probe_duration(audio_file)
                self._set(iid, "解码 + 响度归一化…", None, {"duration": dur}, title=title)
                wav = engine.to_wav16k(audio_file)

                # ---- 阶段一：初稿，逐块推给前端 ----
                def on_chunk(segs):
                    st = PROGRESS.get(iid) or {}
                    st["partial"] = segs
                    PROGRESS[iid] = st

                draft = engine.transcribe_draft(wav, progress=prog, on_chunk=on_chunk)
                if iid in DELETED:
                    print(f"[transcribe] item deleted during draft phase iid={iid}", flush=True)
                    self._set(iid, "已删除", status="error", msg="记录已被删除")
                    return
                if draft:
                    needs_second_pass = auto_diarization or transcription_mode == "accuracy"
                    save_item(iid, {"speakers": {"0": "说话人1"},
                                    "speaker_colors": make_speaker_colors(iid, [0]),
                                    "segments": draft,
                                    "spk_pending": needs_second_pass})
                    with INDEX_LOCK:
                        items = load_index()
                        existing = next((item for item in items if item.get("id") == iid), None)
                        if existing:
                            existing.update({"duration": dur, "n_speakers": 1, "status": "draft"})
                        else:
                            items.insert(0, {"id": iid, "title": title, "duration": dur,
                                             "created": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                                             "n_speakers": 1, "status": "draft"})
                        save_index(items)
                    if not needs_second_pass:
                        with INDEX_LOCK:
                            items = load_index()
                            for item in items:
                                if item.get("id") == iid: item["status"] = "done"
                            save_index(items)
                        self._set(iid, "完成", 1.0, status="done")
                        print(f"[transcribe] speed draft done iid={iid} segs={len(draft)}", flush=True)
                        return
                    pending_stage = ("文稿就绪，说话人分离中…" if auto_diarization
                                     else "文稿就绪，精细校对中…")
                    self._set(iid, pending_stage, None, status="draft")
                    print(f"[transcribe] draft ready iid={iid} segs={len(draft)}", flush=True)

                # ---- 阶段二：完整管线（含声纹分离） ----
                segments = engine.transcribe_full(wav, progress=prog, mode=transcription_mode)
                if not auto_diarization:
                    segments = [{**segment, "spk": 0} for segment in segments]
                n_spk = len({s["spk"] for s in segments}) or 1
                speakers = {str(i): f"说话人{i + 1}" for i in range(n_spk)}

                if iid in DELETED or not os.path.isdir(item_dir(iid)):
                    # 用户在阶段二期间删除了这条记录，直接放弃
                    print(f"[transcribe] item deleted during spk phase iid={iid}", flush=True)
                    self._set(iid, "已删除", status="error", msg="记录已被删除")
                    return
                save_item(iid, {"speakers": speakers,
                                "speaker_colors": make_speaker_colors(iid, range(n_spk)),
                                "segments": segments})

                with INDEX_LOCK:
                    items = load_index()
                    if not any(x["id"] == iid for x in items):
                        items.insert(0, {
                            "id": iid, "title": title, "duration": dur,
                            "created": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "n_speakers": n_spk,
                            "status": "done",
                        })
                    else:
                        for x in items:
                            if x["id"] == iid:
                                x["n_speakers"] = n_spk
                                x["status"] = "done"
                                x.pop("error", None)
                    save_index(items)
                print(f"[transcribe] done iid={iid} segs={len(segments)}", flush=True)
                self._set(iid, "完成", 1.0, status="done")
            except Exception as e:
                import traceback
                traceback.print_exc()
                print(f"[transcribe-error] iid={iid} {e!r}", flush=True)
                self._set(iid, "出错", status="error", msg=str(e))
                with INDEX_LOCK:
                    items = load_index()
                    existing = next((item for item in items if item.get("id") == iid), None)
                    if existing:
                        existing["status"] = "error"
                        existing["error"] = str(e)
                    else:
                        items.insert(0, {"id": iid, "title": title, "duration": 0,
                                         "created": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                                         "n_speakers": 0, "status": "error", "error": str(e)})
                    save_index(items)
            finally:
                if wav:
                    try:
                        os.unlink(wav)
                    except OSError:
                        pass

    # 说话人改名
    def rename_speaker(self, iid, spk_index, new_name):
        data = load_item(iid)
        data.setdefault("speakers", {})[str(spk_index)] = new_name
        save_item(iid, data)
        return True

    # 记录改名
    def rename_item(self, iid, title):
        items = load_index()
        for x in items:
            if x["id"] == iid:
                x["title"] = title
        save_index(items)
        return True

    def bulk_delete_items(self, ids):
        results = []
        for iid in ids or []:
            if any(item.get("id") == iid for item in load_index()):
                result = self.delete_item(iid)
                if result.get("ok"):
                    results.append(iid)
        return {"ok": True, "count": len(results)}

    def bulk_export_items(self, ids, export_format):
        settings = load_settings()
        fmt = str(export_format or "").lower()
        if fmt not in {"txt", "pdf", "docx"}:
            return {"ok": False, "message": "不支持的导出格式"}
        directory = settings["export_directory"]
        try:
            os.makedirs(directory, exist_ok=True)
            exported = 0
            for iid in ids or []:
                payload = export_payload(iid)
                title = safe_filename(payload["title"] or iid)
                output = os.path.join(directory, f"{title}-{iid}.{fmt}")
                if fmt == "txt": write_txt_export(output, payload)
                elif fmt == "pdf": write_pdf_export(output, payload)
                else: write_docx_export(output, payload)
                record_export_path(iid, "document", output)
                exported += 1
            return {"ok": True, "count": exported, "directory": directory}
        except Exception as exc:
            return {"ok": False, "message": str(exc)}

    def update_segment(self, iid, segment_index, text):
        data = load_item(iid)
        segments = data.get("segments", [])
        index = int(segment_index)
        if index < 0 or index >= len(segments):
            return False
        text = str(text or "").strip()
        if not text:
            return False
        segments[index]["text"] = text
        save_item(iid, data)
        return True

    def retry_item(self, iid):
        audio_file = audio_path_of(iid)
        if not audio_file or not os.path.isfile(audio_file):
            return {"ok": False, "message": "找不到原始录音"}
        meta = next((x for x in load_index() if x["id"] == iid), {})
        PROGRESS[iid] = {"stage": "准备重试…", "pct": None, "info": None,
                         "status": "running", "msg": "", "title": meta.get("title", iid), "partial": []}
        DELETED.discard(iid)
        threading.Thread(target=self._run_transcribe,
                         args=(iid, audio_file, meta.get("title", iid)), daemon=True).start()
        return {"ok": True}

    def retry_diarization(self, iid):
        audio_file = audio_path_of(iid)
        if not audio_file or not os.path.isfile(audio_file):
            return {"ok": False, "message": "找不到原始录音"}
        meta = next((x for x in load_index() if x.get("id") == iid), {})
        PROGRESS[iid] = {"stage": "准备重新进行说话人分离…", "pct": None,
                         "info": None, "status": "running", "msg": "", "title": meta.get("title", iid), "partial": []}
        threading.Thread(target=self._run_diarization,
                         args=(iid, audio_file, meta.get("title", iid)), daemon=True).start()
        return {"ok": True}

    def _run_diarization(self, iid, audio_file, title):
        import engine
        wav = None
        with TRANSCRIBE_LOCK:
            try:
                wav = engine.to_wav16k(audio_file)
                settings = load_settings()
                segments = engine.transcribe_full(wav, progress=lambda stage, pct, info=None:
                                                   self._set(iid, stage, pct, info),
                                                  mode=settings["transcription_mode"])
                if not settings["auto_diarization"]:
                    segments = [{**segment, "spk": 0} for segment in segments]
                speakers = {str(i): f"说话人{i + 1}" for i in sorted({s["spk"] for s in segments})} or {"0": "说话人1"}
                data = load_item(iid)
                data.update({"speakers": speakers, "segments": segments, "spk_pending": False})
                save_item(iid, data)
                with INDEX_LOCK:
                    items = load_index()
                    for item in items:
                        if item.get("id") == iid:
                            item["n_speakers"] = len(speakers)
                            item["status"] = "done"
                            item.pop("error", None)
                    save_index(items)
                self._set(iid, "完成", 1.0, status="done")
            except Exception as exc:
                self._set(iid, "说话人分离失败", None, status="error", msg=str(exc))
            finally:
                if wav:
                    try: os.unlink(wav)
                    except OSError: pass

    # 删除
    def delete_item(self, iid):
        DELETED.add(iid)  # 若该条还在转写，阻止后续阶段把它重新落盘
        items = load_index()
        meta = next((item for item in items if item["id"] == iid), {})
        if not load_settings()["delete_audio_with_transcript"]:
            try:
                preserve_item_audio(iid, meta.get("title") or iid)
            except OSError:
                pass
        items = [x for x in items if x["id"] != iid]
        save_index(items)
        source_dir = item_dir(iid)
        if os.path.isdir(source_dir):
            os.makedirs(TRASH, exist_ok=True)
            trash_dir = os.path.join(TRASH, iid)
            shutil.rmtree(trash_dir, ignore_errors=True)
            shutil.move(source_dir, trash_dir)
            with open(os.path.join(trash_dir, "meta.json"), "w", encoding="utf-8") as file:
                json.dump(meta, file, ensure_ascii=False)
        settings = load_settings()
        if settings.get("last_item_id") == iid:
            apply_settings_patch({"last_item_id": ""})
        return {"ok": True, "id": iid, "title": meta.get("title") or iid,
                "audio_preserved": not settings["delete_audio_with_transcript"], "undo": True}

    def restore_deleted_item(self, iid):
        source = os.path.join(TRASH, str(iid))
        target = item_dir(iid)
        if not os.path.isdir(source) or os.path.exists(target):
            return {"ok": False, "message": "这份文稿已无法恢复"}
        shutil.move(source, target)
        transcript = os.path.join(target, "transcript.json")
        if not os.path.isfile(transcript):
            shutil.rmtree(target, ignore_errors=True)
            return {"ok": False, "message": "恢复失败：文稿文件不完整"}
        data = load_item(iid)
        meta = {}
        try:
            with open(os.path.join(target, "meta.json"), encoding="utf-8") as file:
                meta = json.load(file)
            os.remove(os.path.join(target, "meta.json"))
        except (OSError, ValueError):
            pass
        items = load_index()
        if not any(x.get("id") == iid for x in items):
            items.insert(0, {"id": iid, "title": meta.get("title", iid),
                             "duration": meta.get("duration", 0),
                             "created": meta.get("created") or datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                             "n_speakers": meta.get("n_speakers") or len(data.get("speakers", {})) or 1})
            save_index(items)
        return {"ok": True}

    def clear_history(self):
        if TRANSCRIBE_LOCK.locked():
            return {"ok": False, "message": "仍有音频正在转写，请完成后再清理。"}
        keep_audio = not load_settings()["delete_audio_with_transcript"]
        items = load_index()
        for item in items:
            iid = item["id"]
            DELETED.add(iid)
            if keep_audio:
                try:
                    preserve_item_audio(iid, item.get("title") or iid)
                except OSError:
                    pass
            shutil.rmtree(item_dir(iid), ignore_errors=True)
        save_index([])
        apply_settings_patch({"last_item_id": ""})
        return {"ok": True, "count": len(items)}

    def clear_model_cache(self):
        if TRANSCRIBE_LOCK.locked():
            return {"ok": False, "message": "仍有音频正在转写，请完成后再清理模型。"}
        caches = model_cache_dirs()
        freed = sum(path_size(path) for path in caches)
        for path in caches:
            shutil.rmtree(path, ignore_errors=True)
        return {"ok": True, "freed": freed}

    # 导出 Word / PDF / TXT（弹原生保存对话框选位置）
    def export_document(self, iid, export_format=None):
        import webview
        settings = load_settings()
        fmt = export_format if export_format in {"docx", "pdf", "txt"} else settings["export_format"]
        payload = export_payload(iid)
        title = safe_filename(payload["title"])
        if settings["filename_rule"] == "source_date":
            date = (payload["created"] or datetime.datetime.now().strftime("%Y-%m-%d"))[:10].replace("-", "")
            title = f"{title}-{date}"
        default = f"{title}.{fmt}"
        directory = settings["export_directory"]
        res = self.window.create_file_dialog(
            webview.SAVE_DIALOG, save_filename=default,
            directory=directory)
        if not res:
            return None
        out = res if isinstance(res, str) else res[0]
        if os.path.splitext(out)[1].lower() != f".{fmt}":
            out += f".{fmt}"
        if fmt == "docx":
            write_docx_export(out, payload)
        elif fmt == "pdf":
            write_pdf_export(out, payload)
        else:
            write_txt_export(out, payload)
        record_export_path(iid, "document", out)
        return out

    def export_ai_summary(self, iid, export_format=None):
        import webview
        settings = load_settings()
        fmt = export_format if export_format in {"docx", "pdf", "txt"} else settings["export_format"]
        payload = ai_summary_export_payload(iid)
        title = safe_filename(payload["title"])
        if settings["filename_rule"] == "source_date":
            date = (payload["created"] or datetime.datetime.now().strftime("%Y-%m-%d"))[:10].replace("-", "")
            title = f"{title}-{date}"
        default = f"{title}.{fmt}"
        result = self.window.create_file_dialog(
            webview.SAVE_DIALOG, save_filename=default,
            directory=settings["export_directory"])
        if not result:
            return None
        output = result if isinstance(result, str) else result[0]
        if os.path.splitext(output)[1].lower() != f".{fmt}":
            output += f".{fmt}"
        if fmt == "docx": write_ai_docx_export(output, payload)
        elif fmt == "pdf": write_ai_pdf_export(output, payload)
        else: write_txt_export(output, payload)
        record_export_path(iid, "summary", output)
        return output

    def export_txt(self, iid):
        return self.export_document(iid, "txt")

def _preload_model():
    # 启动即后台加载模型（约 30-40 秒），首次转写不用再等
    try:
        import engine
        engine.get_model()
        print("[preload] 模型加载完成", flush=True)
    except Exception as e:
        print(f"[preload] 模型加载失败（转写时会重试）: {e!r}", flush=True)


def _set_dock_icon():
    # 直接用 python 跑时进程会显示 Python 的火箭图标，这里换成本应用的。
    # 必须等主运行循环启动后再设，否则 app 完成启动时会被重置回火箭。
    try:
        import AppKit
        from Foundation import NSOperationQueue
        path = next((p for p in (os.path.join(HERE, "assets", "icon.icns"),
                                 os.path.join(HERE, "icon.icns"))
                     if os.path.exists(p)), None)
        if not path:
            print("[dock] 未找到 icon.icns，跳过", flush=True)
            return

        def apply():
            img = AppKit.NSImage.alloc().initWithContentsOfFile_(path)
            if img:
                application = AppKit.NSApplication.sharedApplication()
                application.setApplicationIconImage_(img)
                # Dock 图标和最小化窗口缩略图是两套状态。只设置前者时，
                # macOS 仍可能从 Python 解释器取回火箭图标。
                for native_window in application.windows():
                    if hasattr(native_window, "setMiniwindowImage_"):
                        native_window.setMiniwindowImage_(img)
                    if hasattr(native_window, "setMiniwindowTitle_"):
                        native_window.setMiniwindowTitle_("WordGrab")
                print("[dock] Dock 与最小化窗口图标已设置", flush=True)
            else:
                print("[dock] icns 加载失败", flush=True)

        NSOperationQueue.mainQueue().addOperationWithBlock_(apply)
    except Exception as e:
        print(f"[dock] 设置图标失败: {e!r}", flush=True)


def _integrate_native_titlebar(window):
    """让网页内容延伸到标题栏，同时保留 macOS 原生红黄绿按钮。

    这不是彻底移除 NSWindow 标题栏，而是把它变成透明的一体式区域。
    因此关闭、最小化、全屏、顶部拖动和双击缩放仍由 macOS 处理。
    """
    try:
        import AppKit

        native = window.native
        if native is None:
            print("[window] 原生窗口尚未就绪，跳过标题栏融合", flush=True)
            return

        full_size = getattr(
            AppKit,
            "NSWindowStyleMaskFullSizeContentView",
            getattr(AppKit, "NSFullSizeContentViewWindowMask", 1 << 15),
        )
        native.setStyleMask_(native.styleMask() | full_size)
        native.setTitlebarAppearsTransparent_(True)
        native.setTitleVisibility_(getattr(AppKit, "NSWindowTitleHidden", 1))
        native.setMovableByWindowBackground_(True)

        icon_path = next((p for p in (os.path.join(HERE, "assets", "icon.icns"),
                                      os.path.join(HERE, "icon.icns"))
                          if os.path.exists(p)), None)
        if icon_path:
            mini_icon = AppKit.NSImage.alloc().initWithContentsOfFile_(icon_path)
            if mini_icon and hasattr(native, "setMiniwindowImage_"):
                native.setMiniwindowImage_(mini_icon)
                native.setMiniwindowTitle_("WordGrab")

        # macOS 11+ 可去掉标题栏与内容之间的系统分隔线。
        if hasattr(native, "setTitlebarSeparatorStyle_"):
            native.setTitlebarSeparatorStyle_(
                getattr(AppKit, "NSTitlebarSeparatorStyleNone", 1)
            )

        # 明确保留系统原生按钮；不在网页里重新绘制一套假的按钮。
        for kind in (
            AppKit.NSWindowCloseButton,
            AppKit.NSWindowMiniaturizeButton,
            AppKit.NSWindowZoomButton,
        ):
            button = native.standardWindowButton_(kind)
            if button is not None:
                button.setHidden_(False)

        print("[window] 已启用一体化原生标题栏", flush=True)
    except Exception as e:
        # 非 macOS 或 PyObjC 不可用时仍可退回普通标题栏启动。
        print(f"[window] 标题栏融合失败: {e!r}", flush=True)


def _install_native_drag_strip(window):
    """在 WebView 顶部覆盖一条透明 NSView，专门接管拖动和双击缩放。

    full-size content view 会让网页覆盖原生标题栏，WKWebView 因而会吞掉鼠标拖动。
    透明拖动带位于最上方 32pt，并从 x=88pt 开始避开红黄绿按钮。
    """
    try:
        import AppKit
        from Foundation import NSNotificationCenter, NSOperationQueue

        global _DRAG_STRIP_CLASS
        if _DRAG_STRIP_CLASS is None:
            class WordGrabDragStrip(AppKit.NSView):
                def acceptsFirstMouse_(self, event):
                    return True

                def mouseDown_(self, event):
                    native = self.window()
                    if native is None:
                        return
                    if event.clickCount() >= 2:
                        native.zoom_(None)
                    else:
                        native.performWindowDragWithEvent_(event)

            _DRAG_STRIP_CLASS = WordGrabDragStrip

        def apply():
            native = window.native
            if native is None or native.contentView() is None:
                print("[window] WebView 尚未就绪，跳过拖动带", flush=True)
                return

            key = id(native)
            center = NSNotificationCenter.defaultCenter()

            # loaded 理论上只触发一次；仍先清理旧监听，避免页面重新载入后重复回调。
            previous_state = _DRAG_OBSERVERS.pop(key, None)
            if previous_state:
                for token in previous_state.get("tokens", ()):
                    center.removeObserver_(token)
                for timer in previous_state.get("timers", ()):
                    try:
                        timer.cancel()
                    except Exception:
                        pass

            def sync_strip():
                """按当前 contentView 的真实坐标重新安装并置顶拖动带。"""
                content = native.contentView()
                if content is None:
                    return

                strip = _DRAG_STRIPS.get(key)
                if strip is None:
                    strip = _DRAG_STRIP_CLASS.alloc().initWithFrame_(
                        AppKit.NSMakeRect(0, 0, 0, 0)
                    )
                    _DRAG_STRIPS[key] = strip
                elif strip.superview() is not None and strip.superview() != content:
                    # 全屏切换等情况下 pywebview 可能替换 contentView。
                    strip.removeFromSuperview()

                bounds = content.bounds()
                drag_height = min(32.0, max(0.0, bounds.size.height))
                drag_left = min(88.0, max(0.0, bounds.size.width))
                is_flipped = bool(content.isFlipped())
                top_y = (bounds.origin.y if is_flipped else
                         bounds.origin.y + max(0.0, bounds.size.height - drag_height))
                frame = AppKit.NSMakeRect(
                    bounds.origin.x + drag_left,
                    top_y,
                    max(0.0, bounds.size.width - drag_left),
                    drag_height,
                )
                strip.setFrame_(frame)
                vertical_mask = (AppKit.NSViewMaxYMargin if is_flipped
                                 else AppKit.NSViewMinYMargin)
                strip.setAutoresizingMask_(AppKit.NSViewWidthSizable | vertical_mask)
                strip.setHidden_(False)

                # WKWebView 启动完成后可能会重新调整自己的子视图层级。
                # 每次同步都重新放到最上层，避免第一次打开时拖动带被网页盖住；
                # 这也是之前“调整高度后才恢复”的根因。
                if strip.superview() == content:
                    strip.removeFromSuperview()
                content.addSubview_positioned_relativeTo_(
                    strip, AppKit.NSWindowAbove, None
                )

            state = {"tokens": [], "callback": None, "pending": False, "timers": []}

            def request_sync(_notification=None):
                # 同一轮缩放可能连续发出多个通知；合并到主线程下一轮更新。
                if state["pending"]:
                    return
                state["pending"] = True

                def update():
                    state["pending"] = False
                    sync_strip()

                NSOperationQueue.mainQueue().addOperationWithBlock_(update)

            state["callback"] = request_sync  # 持有 Python block，避免被回收。
            notification_names = (
                AppKit.NSWindowDidResizeNotification,
                AppKit.NSWindowDidEndLiveResizeNotification,
                AppKit.NSWindowDidEnterFullScreenNotification,
                AppKit.NSWindowDidExitFullScreenNotification,
                AppKit.NSWindowDidChangeScreenNotification,
                AppKit.NSWindowDidChangeBackingPropertiesNotification,
                AppKit.NSWindowDidBecomeKeyNotification,
                AppKit.NSWindowDidBecomeMainNotification,
            )
            for name in notification_names:
                token = center.addObserverForName_object_queue_usingBlock_(
                    name, native, NSOperationQueue.mainQueue(), request_sync
                )
                state["tokens"].append(token)
            _DRAG_OBSERVERS[key] = state

            sync_strip()

            # 首次显示时 WebView 仍可能进行一到两轮布局。仅在 loaded 阶段
            # 安装会导致拖动带被后续布局盖住，直到 resize 通知才重新出现。
            # 在主线程延迟重试几次，确保刚打开就可以拖动。
            for delay in (0.05, 0.18, 0.45, 0.9):
                def retry_sync(_delay=delay):
                    NSOperationQueue.mainQueue().addOperationWithBlock_(sync_strip)

                timer = threading.Timer(delay, retry_sync)
                timer.daemon = True
                timer.start()
                state["timers"].append(timer)

            print("[window] 顶部透明拖动带已启用，并监听窗口尺寸/焦点变化（启动后自动复位）", flush=True)

        NSOperationQueue.mainQueue().addOperationWithBlock_(apply)
    except Exception as e:
        print(f"[window] 拖动带安装失败: {e!r}", flush=True)


def main():
    global API_REF
    import webview

    IS_MACOS = sys.platform == "darwin"

    api = Api()
    API_REF = api

    # Dock 图标设置只在 macOS 执行
    if IS_MACOS:
        _set_dock_icon()

    port = start_server()
    api.port = port
    threading.Thread(target=_preload_model, daemon=True).start()

    # macOS：无边框窗口 + 原生标题栏美化
    # Windows：普通窗口（自带标题栏、关闭按钮、拖动功能）
    window = webview.create_window(
        "WordGrab", url=f"http://127.0.0.1:{port}/",
        js_api=api,
        width=1120, height=740, min_size=(800, 600),
        background_color="#FFFFFF",
        frameless=IS_MACOS,  # 只在 macOS 无边框
        easy_drag=False,
    )
    api.window = window

    # macOS 专属窗口美化（一体式标题栏、自定义拖动条）
    if IS_MACOS:
        window.events.before_show += _integrate_native_titlebar
        window.events.loaded += _install_native_drag_strip
        # loaded 之后 WebView 还会继续完成原生布局；shown 阶段再安装一次，
        # 避免首次打开必须先调整窗口高度才能触发拖动带。
        window.events.shown += _install_native_drag_strip
        window.events.shown += _set_dock_icon

    webview.start(debug=("--debug" in sys.argv))


if __name__ == "__main__":
    main()
